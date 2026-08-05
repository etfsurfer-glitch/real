# -*- coding: utf-8 -*-
"""대전 선도지구 발표 후 매물·호가 동향 — 기자 전달용 데이터 분석 보고서.
수치는 2026-07-24 complex_daily_agg·transactions 실측.
Run: python3 design/press/make_daejeon_seondo_report.py
 → design/press/대전선도지구_발표후동향.docx / .pdf + data/daejeon_seondo_chart.png
발송 없음 — 파일 생성만.

※ 보도자료가 아니라 취재 참고자료(보고서)다.
  - 대표 인용문·홍보 없음. 기자가 직접 검증할 원자료를 표로 준다.
  - '선정되면 오른다'는 식의 예측·단정 금지 — 관측된 사실(매물 잠김)과 그 한계만.
  - 실거래는 신고 30일 지연이라 발표 후 구간은 미완성 — 반드시 명시.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent / "fonts"
DATA = HERE / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "daejeon_seondo_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

DOCTYPE = "데이터 분석 보고서 (취재 참고자료)"
TITLE = "대전 선도지구 발표 후 열흘, 6개 단지 매물 40% 잠겼다"
SUBTITLE = "7월 15일 선정 발표 직후부터 매물이 매일 감소 · 같은 동네 비선정 단지는 –5~–9% · 대전 전체 –2% · 2026년 7월 24일 기준"

SUMMARY = [
    ("분석 대상", "대전 선도지구 6개 단지(크로바·목련·한가람·공작한양·보람·삼익소월) 7,797세대"),
    ("핵심 관측", "발표(7/15) 직후부터 매매 광고매물이 매일 줄어 227건 → 135건(–40.5%)"),
    ("대조군", "같은 동네 비선정 아파트 –5~–9%, 대전 전체 아파트 –2.3%"),
    ("잠김 경로", "신규 등록은 그대로(43→45건), 기존 매물의 회수가 감소를 이끔"),
    ("유의 사항", "실거래는 신고 30일 지연으로 발표 후 구간이 미완성 — 거래량·가격은 다음 달부터 확인"),
]

BODY = [
    ("1. 무엇을 봤나",
     "2026년 7월 15일 대전 노후계획도시 정비 선도지구로 6개 단지가 선정됐습니다(둔산 13구역 목련·크로바, "
     "14구역 한가람·공작한양, 송촌 6구역 보람·삼익소월, 총 7,797세대). 저희는 이 6개 단지의 일자별 "
     "매물수와 호가를 발표 전후로 추적했습니다. 매물 데이터는 저희가 매일 수집하는 것으로, 실거래(국토부 "
     "신고)보다 시장 반응이 즉각 드러납니다. 본 보고서는 관측된 사실과 그 한계를 정리한 것이며, 향후 "
     "가격 방향을 예측하지 않는다."),
    ("2. 매물이 발표 직후부터 매일 줄었습니다",
     "6개 단지의 매매 광고매물은 발표 전날(7/14) 227건에서 7월 24일 135건으로 40.5% 줄었습니다. 같은 집을 "
     "여러 중개사무소가 올린 중복 광고를 걷어낸 실매물 기준으로도 144건에서 99건으로 31.3% 감소했습니다. "
     "특징은 하락이 특정일에 몰린 게 아니라 발표일 이후 열흘간 거의 매일 이어졌다는 점입니다(227→220→"
     "206→195→191→187→176→169→158→142→135). 선정이 곧바로 매물을 거둬들이게 만든 것으로 보입니다."),
    ("3. 같은 동네 비선정 단지와 비교하면 선명합니다",
     "매물 감소가 대전 전체의 흐름인지 확인하기 위해 같은 동(둔산동·탄방동·법동)의 선정되지 않은 아파트와 "
     "대전 전체 아파트를 대조군으로 놓았습니다. 선도지구가 –40.5%인 동안 비선정 단지는 둔산동 –4.9%, 탄방동 "
     "–9.5%, 법동은 오히려 +25.0%(모수가 작아 변동이 큼)였고, 대전 전체 아파트는 –2.3%에 그쳤습니다. "
     "선도지구 6개 단지에서만 매물이 급격히 잠긴 것으로, 발표 외에 이 차이를 설명할 다른 요인은 관측되지 "
     "않았습니다."),
    ("4. 새 매물이 안 나온 게 아니라, 있던 매물을 거둬들였습니다",
     "매물이 줄어드는 경로는 두 가지입니다 — 새 매물이 안 나오거나(신규 등록 감소), 나와 있던 매물을 거둬들이거나"
     "(회수)입니다. 발표 전후 각 열흘간 신규 등록 건수는 43건 → 45건으로 사실상 변화가 없었습니다. 즉 감소는 새 매물이 "
     "끊겨서가 아니라 기존 매물을 회수했기 때문입니다. 집주인들이 선정 뒤 값이 더 오를 것으로 보고 팔 물건을 "
     "거둔 것으로 해석할 여지가 있으나, 회수 사유 자체는 매물 데이터로 확인되지는 않습니다."),
    ("5. 호가는 아직 뚜렷하게 움직이지 않았습니다",
     "매물이 잠기는 동안에도 6개 단지 평균 매매 호가는 7월 14일 7.68억원에서 7월 24일 7.61억원으로 큰 변화가 "
     "없었습니다. 단지별로는 목련이 13.26억→15.14억으로 올랐으나 이는 남은 매물이 대형 위주로 바뀐 구성효과가 "
     "섞여 있고(매물 15건→7건), 나머지 단지는 대체로 보합이었습니다. 매물이 줄어든 상태에서 산출된 평균이라 "
     "표본이 얇아진 점도 함께 감안해 주시기 바랍니다."),
    ("6. 이 자료로 말할 수 있는 것과 없는 것",
     "말씀드릴 수 있는 것은 ‘발표 직후 선도지구 매물이 다른 지역보다 뚜렷하게 빠르게 잠겼다’ 하나입니다. "
     "말씀드리기 어려운 것은 앞으로의 가격입니다. 실거래는 계약 후 30일 내 신고제라 발표 이후 계약분은 아직 대부분 "
     "신고되지 않았습니다(6개 단지 7월 신고 거래는 10건으로 미완성). 따라서 이번 발표가 실제 거래가격에 어떤 "
     "영향을 줬는지는 8월 이후에야 확인됩니다. 또한 관측 기간이 발표 후 열흘로 짧아, 이 매물 잠김이 일시적 관망인지 "
     "지속될 흐름인지는 더 지켜봐야 합니다."),
]

# 원자료 — 단지별 (기자 검증용)
RAW = [
    ("크로바", "서구 둔산동", 1632, 54, 31, 28, 21, "17.6억", "17.6억"),
    ("공작한양", "서구 탄방동", 1074, 53, 31, 31, 22, "5.98억", "6.08억"),
    ("삼익소월", "대덕구 법동", 810, 36, 26, 25, 19, "1.84억", "1.84억"),
    ("한가람", "서구 탄방동", 1380, 31, 13, 25, 11, "4.66억", "4.50억"),
    ("보람", "대덕구 법동", 1735, 30, 27, 20, 19, "3.34억", "3.16억"),
    ("목련", "서구 둔산동", 1166, 23, 7, 15, 7, "13.3억", "15.1억"),
]

# 일자별 매물 추이 (7/10~7/24, 매매 광고)
DAILY = [
    ("07-10", 231), ("07-11", 232), ("07-12", 223), ("07-13", 231), ("07-14", 227),
    ("07-15", 220), ("07-16", 206), ("07-17", 195), ("07-18", 191), ("07-19", 187),
    ("07-20", 176), ("07-21", 169), ("07-22", 158), ("07-23", 142), ("07-24", 135),
]

# 대조군 (7/14 → 7/24 매매 광고 변화율 %)
CONTROL = [
    ("선도지구 6개 단지", -40.5),
    ("탄방동 (비선정)", -9.5),
    ("둔산동 (비선정)", -4.9),
    ("대전 전체 아파트", -2.3),
    ("법동 (비선정)", 25.0),
]

NOTE = (
    "작성 | 콕집(koczip.com) 데이터팀 · 런투온라인(대표 황인찬) · runtoonline@gmail.com · 010-5942-8014  "
    "모든 수치는 저희가 자체 구축한 DB에서 2026년 7월 24일 기준으로 집계했습니다. 매물수는 콕집 일별 수집 "
    "기준으로, 광고매물(포털 노출 광고 건수)과 실매물(같은 집의 중복 광고를 합친 수)을 구분해 표기했습니다. "
    "선도지구 6개 단지는 국토부 선정 발표문 기준입니다. 실거래는 국토부 신고 자료(해제거래 제외)이며 신고 "
    "기한 30일로 최근 구간은 미완성입니다. 대조군은 같은 법정동의 선정되지 않은 아파트입니다. "
    "분석 방법이나 원자료가 더 필요하시면 언제든 연락 주시기 바랍니다."
)


def make_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10.8, 3.9), dpi=200)

    # ① 일자별 매물 추이 — 발표일(7/15) 세로선
    xs = list(range(len(DAILY)))
    ys = [d[1] for d in DAILY]
    ax1.plot(xs, ys, color="#d64545", lw=2.2, marker="o", ms=4)
    ann_i = 5  # 7/15
    ax1.axvline(ann_i, color="#334155", lw=1.1, ls="--")
    ax1.annotate("7/15 선정 발표", (ann_i, max(ys)), fontsize=8.5, fontweight="bold",
                 color="#334155", xytext=(ann_i + 0.2, max(ys) - 2), va="top")
    ax1.annotate(f"{ys[0]}건", (0, ys[0]), fontsize=9, fontweight="bold", color="#d64545",
                 xytext=(0, 6), textcoords="offset points", ha="left")
    ax1.annotate(f"{ys[-1]}건", (xs[-1], ys[-1]), fontsize=9, fontweight="bold", color="#d64545",
                 xytext=(0, -14), textcoords="offset points", ha="right")
    ax1.set_xticks(xs[::2])
    ax1.set_xticklabels([DAILY[i][0] for i in xs[::2]], fontsize=8)
    ax1.set_ylabel("매매 광고매물 (건)", fontsize=9)
    ax1.set_ylim(120, 245)
    ax1.set_title("① 발표 직후부터 매일 감소 — 6개 단지 합계",
                  fontsize=10.5, fontweight="bold", pad=8, loc="left")
    ax1.grid(axis="y", color="#eef2f6", lw=.8)
    for sp in ("top", "right"):
        ax1.spines[sp].set_visible(False)

    # ② 대조군 변화율
    names = [c[0] for c in CONTROL]
    vals = [c[1] for c in CONTROL]
    y = list(range(len(names)))[::-1]
    cols = ["#d64545" if v < -20 else ("#94a3b8" if v <= 0 else "#4a9d6a") for v in vals]
    ax2.barh(y, vals, 0.62, color=cols)
    ax2.axvline(0, color="#333", lw=1.0)
    for yy, v in zip(y, vals):
        ax2.annotate(f"{v:+.1f}%", (v, yy), va="center",
                     ha="left" if v >= 0 else "right",
                     xytext=(4 if v >= 0 else -4, 0), textcoords="offset points",
                     fontsize=9, fontweight="bold")
    ax2.set_yticks(y)
    ax2.set_yticklabels(names, fontsize=9)
    ax2.set_xlim(-52, 40)
    ax2.set_xlabel("매매 매물 변화율 (7/14 → 7/24, %)", fontsize=9)
    ax2.set_title("② 선도지구만 유독 빠르게 잠겼다",
                  fontsize=10.5, fontweight="bold", pad=8, loc="left")
    ax2.tick_params(axis="x", labelsize=8)
    for sp in ("top", "right", "left"):
        ax2.spines[sp].set_visible(False)

    fig.suptitle("대전 선도지구 발표 전후 매물 동향 (콕집 DB, 2026-07-24)",
                 fontsize=11.3, fontweight="bold", y=0.99)
    fig.text(0.5, 0.005,
             "매물수는 콕집 일별 수집 기준(광고매물). 법동(비선정)의 +25%는 매물 16건의 작은 모수에서 나온 변동입니다. "
             "실거래가 아니라 매물(호가) 데이터이며, 발표 후 실거래는 신고 지연으로 아직 미완성입니다.",
             ha="center", fontsize=7.8, color="#64748b")
    fig.tight_layout(rect=[0, 0.045, 1, 0.94])
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"; st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45; st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run(DOCTYPE)
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    t = d.add_paragraph(); r = t.add_run(TITLE); r.font.size = Pt(14.5); r.font.bold = True
    s2 = d.add_paragraph(); r = s2.add_run(SUBTITLE)
    r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 요약"); r.font.bold = True; r.font.size = Pt(11.5)
    tb = d.add_table(rows=0, cols=2); tb.style = "Light Grid Accent 1"
    for k, v in SUMMARY:
        c = tb.add_row().cells
        c[0].paragraphs[0].add_run(k).bold = True
        c[1].paragraphs[0].add_run(v)
    for rowx in tb.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.6))

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)

    h = d.add_paragraph(); r = h.add_run("■ 부록 — 단지별 매물·호가 (7/14 발표 전날 → 7/24 현재)")
    r.font.bold = True; r.font.size = Pt(11.5)
    HEADS = ["단지", "소재지", "세대", "광고매물\n7/14", "광고매물\n7/24",
             "실매물\n7/14", "실매물\n7/24", "평균호가\n7/14", "평균호가\n7/24"]
    tb2 = d.add_table(rows=1, cols=len(HEADS)); tb2.style = "Light Grid Accent 1"
    for c, x in zip(tb2.rows[0].cells, HEADS):
        c.paragraphs[0].add_run(x).bold = True
    for row in RAW:
        cells = tb2.add_row().cells
        vals = [row[0], row[1], f"{row[2]:,}", str(row[3]), str(row[4]),
                str(row[5]), str(row[6]), row[7], row[8]]
        for c, x in zip(cells, vals):
            c.paragraphs[0].add_run(x)
    for rowx in tb2.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(8)

    m = d.add_paragraph(); r = m.add_run(NOTE)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    d.save(out); print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE); gray = colors.Color(0.39, 0.45, 0.55)
    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=3),
        "title": ParagraphStyle("title", fontName="PB", fontSize=14, leading=19, spaceAfter=3),
        "sub": ParagraphStyle("sub", fontName="P", fontSize=9.4, textColor=gray, leading=14, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "note": ParagraphStyle("note", fontName="P", fontSize=7.6, leading=11.5, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=7.4, leading=10),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=7.4, leading=10),
        "sumk": ParagraphStyle("sumk", fontName="PB", fontSize=8.8, leading=12.5),
        "sumv": ParagraphStyle("sumv", fontName="P", fontSize=8.8, leading=12.5),
    }
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=15 * mm, rightMargin=15 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart = RLImage(str(CHART), width=180 * mm, height=180 * mm * h / w)

    sum_tbl = Table([[Paragraph(k, ss["sumk"]), Paragraph(v, ss["sumv"])] for k, v in SUMMARY],
                    colWidths=[24 * mm, 156 * mm])
    sum_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.955, 0.97, 0.99)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))

    HEADS = ["단지", "소재지", "세대", "광고매물\n7/14", "광고매물\n7/24",
             "실매물\n7/14", "실매물\n7/24", "평균호가\n7/14", "평균호가\n7/24"]
    tdata = [[Paragraph(x, ss["cellb"]) for x in HEADS]]
    for row in RAW:
        vals = [row[0], row[1], f"{row[2]:,}", str(row[3]), str(row[4]),
                str(row[5]), str(row[6]), row[7], row[8]]
        tdata.append([Paragraph(x, ss["cell"]) for x in vals])
    raw_tbl = Table(tdata, colWidths=[20 * mm, 22 * mm, 14 * mm, 14 * mm, 14 * mm,
                                      16 * mm, 16 * mm, 22 * mm, 22 * mm])
    raw_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))

    story = [
        Paragraph(DOCTYPE, ss["tag"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph("■ 요약", ss["head"]), sum_tbl, Spacer(1, 4 * mm),
        chart, Spacer(1, 3 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph("■ 부록 — 단지별 매물·호가 (7/14 발표 전날 → 7/24 현재)", ss["head"]))
    story.append(raw_tbl)
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(NOTE, ss["note"]))
    doc.build(story); print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "대전선도지구_발표후동향.docx")
    make_pdf(HERE / "대전선도지구_발표후동향.pdf")
