# -*- coding: utf-8 -*-
"""반도체 벨트 8개 도시 — 인구이동 · 실거래 · 공급 교차분석 보도자료.
수치는 2026-07-24 pop_move(행안부 인구이동) · transactions · complexes 실측.
※ 현행화(2026-07-24): 인구·공급률은 창 고정으로 불변, 실거래만 30일 신고 누적분 반영
  (26Q2 거래건수 소폭 증가, 변동률 아산 4.7→4.6·구미 8.8→8.9·용인 5.9→6.0·이천 -5.2→-5.3,
   2026 ㎡당 평택 430→431·이천 308→306·청주 377→376).
  변동률은 정밀 ㎡ 평균비율 기준(표시 만원값은 반올림) — 원 방법론 유지.
※ 검증(2026-07-24): 본문 2건 정정 — "순유입 상위 2곳=가격 하위 2곳"은 오류(화성은 가격 1위)라
  "순유입 1·6위(평택·이천)가 가격 7·8위"로, "상승률 높은 청주·아산"에서 아산(+4.6%, 6위)은
  공급 가설 반례로 정직하게 표기.
Run: python3 design/press/make_semi_belt.py
 → design/press/반도체벨트_인구와집값.docx / .pdf + data/semi_belt_chart.png
발송 없음 — 파일 생성만.

※ 당초 가설('양질의 일자리 → 인구유입 → 집값 상승')은 이 8개 지역에서 성립하지 않았다.
   순유입 1위 평택이 가격 상승 7위, 순유입 최하위 구미·용인이 중위권이다.
   데이터가 지지하지 않는 인과는 쓰지 않는다. 대신 '갈렸다'는 사실과, 공급률이라는
   설명 후보를 제시하되 표본 8개로 법칙을 주장하지 않는다.
※ 인구이동은 2024년 7월부터 API가 구 단위 제공을 중단하고 시 단위로 합산한다.
   따라서 전 지표를 '시 단위'로 통일했다. 동탄구는 화성시에 합산되어 별도 산출 불가.
※ 화성시 가격 +19.1%에는 구성효과가 섞여 있다(동탄 거래비중 56%→70%).
   동탄구 자체는 +9.1%. 본문·주석에 반드시 병기한다.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent / "fonts"
DATA = HERE / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "semi_belt_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

TITLE = "인구 유입 1위 평택, 집값은 3년째 제자리 — 반도체 벨트서 인구와 집값이 갈렸다"
SUBTITLE = "콕집, 반도체 8개 도시 인구이동·실거래·공급 교차분석 — 순유입 1·6위가 가격은 7·8위 · 두 곳 모두 공급률 1·2위"

LEAD = (
    "반도체 산업단지를 낀 지역에 인구가 몰리면 집값이 오른다는 통념이 실제 데이터에서는 "
    "확인되지 않았다. 부동산 데이터 플랫폼 콕집(koczip.com)이 반도체 관련 8개 도시의 최근 1년 "
    "인구 순유입과 아파트 실거래가를 함께 분석한 결과, 순유입이 가장 많은 평택(1만5,638명)의 "
    "㎡당 실거래가 상승률은 2.6%로 8곳 중 7위에 그쳤다. 반대로 인구가 순유출된 구미(-1,460명)와 "
    "용인(-2,532명)은 각각 8.9%, 6.0% 올라 중위권이었다. 평택은 순유입이 1년 만에 2.2배로 늘어나는 "
    "동안 ㎡당 가격이 2023년 435만원에서 2026년 431만원으로 오히려 낮아져, 3년째 제자리걸음이다. "
    "인구 유입과 가격 상승의 순위가 서로 어긋난 가운데, 두 지표 사이에서 뚜렷하게 갈린 것은 "
    "신규 공급 물량이었다."
)

BODY = [
    ("① 인구는 분명히 늘었다 — 평택 2.2배, 청년이 절반",
     "행정안전부 주민등록 인구이동 자료로 최근 1년(2025년 7월~2026년 6월) 순유입을 집계한 결과 "
     "평택이 1만5,638명으로 8개 도시 중 가장 많았다. 직전 1년(7,077명)의 2.2배다. 이 가운데 20·30대가 "
     "7,887명으로 절반(50.4%)을 차지해, 젊은 층이 실제로 유입되고 있다는 점도 확인됐다. 화성시가 "
     "1만3,658명으로 뒤를 이었고 20·30대는 8,892명(65.1%)이었다. 아산 5,883명, 천안 3,625명, "
     "청주 3,476명 순이었다. 반면 구미는 1,460명, 용인은 2,532명이 순유출됐다. 용인은 직전 1년 "
     "9,197명 순유입에서 순유출로 방향이 바뀌었다."),
    ("② 그런데 집값 순위는 인구 순위와 맞지 않았다",
     "같은 8개 도시의 ㎡당 아파트 실거래가를 2025년 2분기와 2026년 2분기로 비교하자 순위가 크게 "
     "어긋났다. 순유입 1위 평택은 421만원에서 432만원으로 2.6% 오르는 데 그쳐 7위였다. 순유입 6위 "
     "이천은 324만원에서 307만원으로 5.3% 내려 유일하게 하락했다. 반대로 인구가 빠져나간 구미는 "
     "234만원에서 254만원으로 8.9%, 용인은 795만원에서 842만원으로 6.0% 올랐다. 상승률이 가장 높은 "
     "곳은 화성시(19.1%)와 청주(11.2%)였다. 순유입 1위 평택과 6위 이천이 가격에서는 하위 두 곳(7·8위)에 "
     "자리한 셈이다. "
     "거래량 역시 인구와 따로 움직였다. 인구가 순유출된 용인의 거래는 4,509건에서 5,140건으로 "
     "14.0% 늘었고, 인구가 늘어난 청주의 거래는 3,591건에서 3,139건으로 12.6% 줄었다."),
    ("③ 갈림길은 공급이었다 — 가격 하위 두 곳이 공급률 1·2위",
     "두 지표가 어긋난 자리에서 설명 후보로 남은 것은 신규 공급이다. 각 도시의 기존 아파트 재고 대비 "
     "2024~2025년 사용승인(준공) 세대수 비율을 계산한 결과, 평택이 9.38%로 가장 높았고 이천이 8.01%로 "
     "그 뒤를 이었다. 가격이 가장 부진했던 두 곳이 공급률에서는 1·2위였다. 평택은 2년간 1만6,667세대, "
     "이천은 3,534세대가 새로 준공됐다. 반대로 상승률이 11.2%로 높았던 청주는 공급률이 5.36%로 낮은 "
     "축이었다. 다만 공급률이 가장 낮은 아산(4.96%)은 상승률도 4.6%에 그쳐, 공급만으로 가격이 "
     "설명되지는 않았다. 인구가 들어와도 그만큼, 혹은 그 이상으로 새 아파트가 공급되면 가격이 눌리는 "
     "경향은 보이되, 8개 도시라는 적은 표본에서 관찰된 대응 관계이므로 공급률이 가격을 결정한다고 "
     "단정하기는 어렵다."),
    ("④ 평택은 기저효과도 아니다 — 3년 내내 눌렸다",
     "평택의 부진이 앞서 급등한 데 따른 반작용인지 확인하기 위해 연도별 ㎡당 실거래가를 살펴봤다. "
     "평택은 2023년 435만원에서 2024년 422만원, 2025년 418만원으로 내린 뒤 2026년 431만원으로 "
     "일부 회복했으나 여전히 2023년 수준에 못 미친다. 3년간 상승 구간 자체가 없었다. 이천도 "
     "335만원에서 306만원으로 3년 연속 내렸다. 같은 기간 동탄구는 843만원에서 996만원으로, "
     "청주는 317만원에서 376만원으로 꾸준히 올랐다. 평택의 정체는 특정 시점의 기저효과가 아니라 "
     "3년에 걸쳐 이어진 흐름이다."),
    ("⑤ 동탄 — 오른 것은 가격보다 거래량이었다",
     "성과급 이슈로 주목받은 동탄의 경우, 화성시 전체 ㎡당 가격은 732만원에서 872만원으로 19.1% "
     "올랐으나 이 수치는 그대로 읽기 어렵다. 화성시 아파트 거래에서 동탄구가 차지하는 비중이 "
     "56%에서 70%로 커졌는데, 동탄이 시내에서 가격대가 높은 지역이어서 시 평균이 끌려 올라간 "
     "구성효과가 섞여 있기 때문이다. 동탄구만 따로 보면 916만원에서 999만원으로 9.1% 올라 "
     "청주(11.2%)보다 낮다. 대신 뚜렷하게 달라진 것은 거래량이다. 동탄구 2분기 거래는 1,798건에서 "
     "4,130건으로 130% 늘었다. 동탄에서 일어난 변화는 가격 급등보다 거래 폭증에 가깝다."),
    ("지역별 실거래·매물, 누구나 확인 가능",
     "콕집은 전국 아파트·오피스텔 6만 4천여 단지의 매물과 실거래를 매일 수집해 교차 분석하는 "
     "플랫폼으로, 이번 분석에 쓰인 시군구·읍면동 단위 실거래 통계와 단지별 매물·호가 추이는 "
     "콕집에서 누구나 무료로 확인할 수 있다. 지역 비교, 부동산 타임머신(정책 연대기), 급매 탐지, "
     "단지별 신고가 등 분석 기능도 함께 제공한다."),
]

QUOTE = ("조용호 콕집 공동창업자는 “반도체 일자리가 늘면 사람이 모이고 집값이 오른다는 순서로 "
         "이야기되지만, 데이터에서는 인구와 가격의 순위가 맞지 않았다”며 “평택은 최근 1년 순유입이 "
         "8개 도시 중 가장 많았고 그 절반이 20·30대였는데도 ㎡당 가격은 3년 전보다 낮다. 인구가 "
         "들어오는 것과 값이 오르는 것은 다른 문제이고, 그 사이에 공급이라는 변수가 있다는 점을 "
         "이번 수치가 보여준다”고 말했다.")

COMPANY = [
    ("서비스", "콕집 — 부동산 매물·실거래·중개사 분석 플랫폼 (koczip.com)"),
    ("운영사", "런투온라인 (대표 황인찬)"),
    ("문의", "조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946"),
]

METHOD = (
    "분석 방법 | 대상: 반도체 생산·소재 거점 8개 도시(평택·화성·용인·이천·청주·천안·아산·구미). "
    "인구는 행정안전부 주민등록 인구이동 자료(공공데이터포털) 2022년 10월~2026년 6월 45개월을 자체 수집, "
    "순유입=전입−전출, 청년=20·30대. 이 자료는 2024년 7월부터 구 단위 제공이 중단돼 시 단위로 합산되므로 "
    "전 지표를 시 단위로 통일했다(동탄구는 화성시에 합산되어 별도 인구 산출 불가). "
    "가격은 국토교통부 실거래 신고 자료(해제거래 제외)의 ㎡당 단가 평균이며, 신고 기한 30일로 최근 2개월은 미완성. "
    "화성시 상승률 19.1%에는 동탄구 거래비중 확대(56%→70%)에 따른 구성효과가 포함되며 동탄구 자체는 +9.1%. "
    "공급률은 콕집 DB의 아파트 단지 사용승인일·세대수 기준 2024~2025년 준공 세대수 ÷ 기존 총세대수로, "
    "포털 미등록 신규 단지가 있을 경우 실제보다 낮게 잡힐 수 있다. "
    "대상이 8개 도시로 적어 지표 간 대응 관계는 경향으로만 제시하며 통계적 유의성은 주장하지 않는다. "
    "전 수치는 콕집 자체 구축 DB 실측(2026-07-24)."
)

# ── 실측 데이터(2026-07-21 DB) ──
# (지역, 순유입, 청년순유입, 가격변동%, 공급률%, 거래25Q2, 거래26Q2)  — 순유입 내림차순
ROWS = [
    ("평택",     15638, 7887,   2.6, 9.38, 1832, 2138),
    ("화성시",   13658, 8892,  19.1, 6.67, 3207, 5886),
    ("아산",      5883, 1978,   4.6, 4.96, 1483, 1270),
    ("천안",      3625, 1782,   5.3, 6.09, 2653, 2370),
    ("청주",      3476, 3113,  11.2, 5.36, 3591, 3139),
    ("이천",      1497,  310,  -5.3, 8.01,  379,  541),
    ("구미",     -1460, -1144,  8.9, 5.52, 1347, 1222),
    ("용인",     -2532,  -510,  6.0, 5.48, 4509, 5140),
]

# 연도별 ㎡당 만원 (기저효과 확인용)
YEARLY = {
    "평택":   [435, 422, 418, 431],
    "이천":   [335, 333, 319, 306],
    "동탄구": [843, 871, 932, 996],
    "청주":   [317, 321, 346, 376],
}

TABLE = [
    (nm, f"{inn:+,}명", f"{yng:+,}명", f"{pr:+.1f}%", f"{sp:.2f}%",
     f"{t25:,}→{t26:,}건 ({100.0*(t26-t25)/t25:+.0f}%)")
    for nm, inn, yng, pr, sp, t25, t26 in ROWS
]


def make_chart():
    """같은 지역 순서(순유입 내림차순)로 세 지표를 나란히 놓는다.
    ①이 계단처럼 내려가는데 ②가 뒤죽박죽이면 '둘이 따로 논다'가 눈으로 보인다.
    ③에서 평택·이천만 튀어나오면 '갈림길은 공급'이 같은 그림에서 읽힌다.
    지표별로 단위가 달라 한 축에 못 얹으므로 패널을 나누되 순서를 고정하는 것이 핵심.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    names = [r[0] for r in ROWS]
    y = list(range(len(names)))[::-1]          # 위에서부터 순유입 큰 순
    fig, axes = plt.subplots(1, 3, figsize=(11.6, 4.3), dpi=200)

    panels = [
        (0, [r[1] for r in ROWS], "① 인구 순유입 (최근 1년)", "명",
         lambda v: f"{v:+,.0f}", "#1268d3"),
        (1, [r[3] for r in ROWS], "② ㎡당 실거래가 변동 (25·2Q→26·2Q)", "%",
         lambda v: f"{v:+.1f}%", "#e08a1e"),
        (2, [r[4] for r in ROWS], "③ 신규 공급률 (재고 대비 24~25년 준공)", "%",
         lambda v: f"{v:.2f}%", "#12a06a"),
    ]
    for idx, vals, title, unit, fmt, base in panels:
        ax = axes[idx]
        # 부진/과공급 지역을 빨강으로 집어 세 패널이 같은 지역을 가리키게 한다
        cols = []
        for nm, v in zip(names, vals):
            hot = (idx == 1 and v < 3) or (idx == 2 and v > 7.5) or (idx == 0 and nm == "평택")
            cols.append("#d64545" if hot else base)
        ax.barh(y, vals, 0.66, color=cols, alpha=.9)
        ax.axvline(0, color="#333", lw=1.0)
        for yy, v in zip(y, vals):
            ax.annotate(fmt(v), (v, yy), va="center",
                        ha="left" if v >= 0 else "right",
                        xytext=(4 if v >= 0 else -4, 0), textcoords="offset points",
                        fontsize=8.2, fontweight="bold")
        ax.set_yticks(y)
        ax.set_yticklabels(names if idx == 0 else [""] * len(names), fontsize=9.5)
        if idx:                       # 지역명은 왼쪽 패널에만 — 빈 눈금선까지 지운다
            ax.tick_params(axis="y", length=0)
        ax.set_xlabel(unit, fontsize=8.5)
        ax.set_title(title, fontsize=10, fontweight="bold", pad=8, loc="left")
        ax.tick_params(axis="x", labelsize=8)
        ax.grid(axis="x", color="#eef2f6", lw=.8)
        for sp in ("top", "right", "left"):
            ax.spines[sp].set_visible(False)
        lo, hi = min(vals + [0]), max(vals + [0])
        pad = (hi - lo) * 0.22
        # 값이 전부 양수인 지표(공급률)까지 축을 음수로 늘리면 '마이너스 공급'처럼 읽힌다
        ax.set_xlim(lo - pad if lo < 0 else 0, hi + pad)

    fig.suptitle("반도체 8개 도시 — 세 패널 모두 '인구 순유입이 많은 순'으로 정렬 (콕집 DB)",
                 fontsize=11.3, fontweight="bold", y=0.985)
    fig.text(0.5, 0.005,
             "①은 위에서부터 순서대로 줄어들지만 ②는 순서를 따르지 않는다 — 인구와 집값이 따로 움직였다는 뜻. "
             "②에서 부진한 평택·이천이 ③에서는 공급률 1·2위로 튄다.",
             ha="center", fontsize=8, color="#64748b")
    fig.tight_layout(rect=[0, 0.045, 1, 0.945])
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 7월 · 즉시 보도 가능    문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(15); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.4))

    HEADS = ["지역", "순유입", "20·30대", "㎡당 가격", "공급률", "2분기 거래량"]
    h = d.add_paragraph(); r = h.add_run("■ 8개 도시 한눈에 비교 (콕집 DB, 순유입 순)")
    r.font.bold = True; r.font.size = Pt(11.5)
    tbl = d.add_table(rows=1, cols=len(HEADS))
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, HEADS):
        c.paragraphs[0].add_run(txt).bold = True
    for row in TABLE:
        for c, txt in zip(tbl.add_row().cells, row):
            c.paragraphs[0].add_run(txt)
    for rowx in tbl.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(8.5)

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    m = d.add_paragraph(); r = m.add_run(METHOD)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=14.5, leading=20, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=10.4, textColor=blue, leading=15, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "method": ParagraphStyle("method", fontName="P", fontSize=7.6, leading=11, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=8, leading=11),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=8, leading=11),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=15 * mm, rightMargin=15 * mm,
                            topMargin=15 * mm, bottomMargin=15 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 180 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    def cell(txt, bold=False):
        return Paragraph(txt, ss["cellb"] if bold else ss["cell"])

    HEADS = ["지역", "순유입", "20·30대", "㎡당 가격", "공급률", "2분기 거래량"]
    tdata = [[cell(x, True) for x in HEADS]]
    for row in TABLE:
        tdata.append([cell(row[0], True)] + [cell(x) for x in row[1:]])
    tbl = Table(tdata, colWidths=[20 * mm, 27 * mm, 25 * mm, 24 * mm, 22 * mm, 42 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))

    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 7월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
        chart,
        Spacer(1, 4 * mm),
        Paragraph("■ 8개 도시 한눈에 비교 (콕집 DB, 순유입 순)", ss["head"]),
        tbl,
        Spacer(1, 2 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph(METHOD, ss["method"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "반도체벨트_인구와집값.docx")
    make_pdf(HERE / "반도체벨트_인구와집값.pdf")
