# -*- coding: utf-8 -*-
"""노후계획도시 선도지구 3개 도시(분당·부산·대전) 궤적 비교 보도자료.
수치는 2026-07-20 transactions·complex_daily_agg 실측.
Run: python3 design/press/make_seondo_3city.py
 → design/press/선도지구_3도시비교.docx / .pdf + data/seondo_3city_chart.png
발송 없음 — 파일 생성만.

※ 단지 특정 근거: 분당은 이름·세대수 역추적(시범 우성+현대=3,569세대, 양지 6단지=4,392세대가
   선정 발표 수치와 일치). 부산은 5단지 합 7,318세대가 발표 수치와 정확히 일치.
   **보도 전 국토부 선정 발표문과 최종 대조 권장**(단지 오기는 신뢰도 직결).
※ '대전도 분당처럼 오른다'는 식의 예측은 쓰지 않는다 — 실제로 부산은 거래가 40% 줄어
   분당과 반대로 갔다. 같은 제도라도 결과가 갈렸다는 것이 데이터가 말하는 바다.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent / "fonts"
DATA = HERE / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "seondo_3city_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

TITLE = "선도지구 선정되면 거래부터 멈춘다 — 분당·부산, 직후 3개월 거래 최대 71% 감소"
SUBTITLE = "콕집, 선정 시점 맞춰 비교 — 가격은 곧바로 7~18% 올라 유지 · 대전은 지금 그 초입"

LEAD = (
    "노후계획도시 정비 선도지구로 선정된 지역은 발표 직후 거래가 먼저 멈추고, 가격은 곧바로 올라 "
    "유지되는 흐름을 보인 것으로 나타났다. 부동산 데이터 플랫폼 콕집(koczip.com)이 앞서 선정된 "
    "분당(2024년 11월, 12개 단지 10,738세대)과 부산(2025년 12월, 5개 단지 7,318세대)을 선정 시점에 "
    "맞춰 비교한 결과다. 선정 후 3개월 안에 매매 거래량은 분당이 최대 71%, 부산이 최대 55% 줄었다. "
    "반면 평균 실거래가는 두 지역 모두 선정 직후부터 올라 분당 7~18%, 부산 9~11% 높은 수준을 "
    "유지했다. 7월 15일 선정된 대전은 발표 엿새째로, 앞선 두 지역이 지난 구간의 초입에 있다."
)

BODY = [
    ("비교 방법 — 선정 시점을 맞추고, 기준선을 세 가지로 검증했다",
     "세 지역은 선정 시점이 각각 2024년 11월(분당), 2025년 12월(부산), 2026년 7월(대전)로 다르다. "
     "경과 기간이 다른 상태를 그대로 나란히 놓으면 비교가 성립하지 않으므로, 각 지역의 선정일을 "
     "기준점으로 삼아 ‘선정 후 1개월·2개월·3개월’ 시점을 같은 자리끼리 맞춰 비교했다. "
     "비교의 기준값은 선정 직전 기간의 월평균 거래량과 평균 실거래가다. 다만 이 단지들은 월 거래가 "
     "20~30건대로 많지 않아 기준 기간을 어떻게 잡느냐에 따라 수치가 달라질 수 있다. 이에 기준 기간을 "
     "선정 직전 3개월·6개월·12개월 세 가지로 각각 잡아 모두 계산했고, 본문에는 세 결과의 범위를 "
     "제시했다. 세 기준에서 모두 같은 방향이 나온 항목만 결론으로 삼았다."),
    ("① 거래 — 선정 직후 3개월간 뚜렷하게 줄었다",
     "분당은 선정 후 1개월 시점 거래량이 선정 전 대비 37~59% 줄었고, 2개월 시점에는 56~71% 감소해 "
     "선정 전의 3분의 1 수준까지 내려갔다. 3개월 시점에도 37~59% 적은 상태가 이어졌다. 부산은 "
     "1개월 시점 10~44%, 2개월 시점 26~54%, 3개월 시점 27~55% 줄었다. 기준 기간을 어떻게 잡아도 "
     "두 지역 모두 세 시점 전부에서 감소가 확인됐다. 선정이 곧바로 거래 증가로 이어지지는 않았다는 "
     "의미다."),
    ("② 가격 — 거래가 멈춘 동안에도 곧바로 올라 유지됐다",
     "거래가 줄어드는 동안에도 평균 실거래가는 두 지역 모두 선정 직후부터 상승했다. 분당은 1개월 "
     "시점 이미 선정 전보다 8~16% 높았고, 3개월 시점에는 11~18% 높은 수준이었다. 부산은 1개월 시점 "
     "10~11%, 3개월 시점 9~10% 높았다. 거래량과 마찬가지로 기준 기간을 세 가지로 바꿔도 두 지역, "
     "세 시점 모두에서 상승 방향이 유지됐다. 다만 거래 건수가 줄어든 상태에서 산출된 평균가라는 "
     "점은 함께 감안해야 한다."),
    ("③ 대전 — 지금은 매물이 잠긴 초입",
     "대전 선도지구 6개 단지는 발표 전날인 7월 14일 227건이던 매매 광고매물이 7월 20일 176건으로 "
     "22.5% 줄었다. 실매물 기준으로도 144건에서 117건으로 18.8% 감소했다. 같은 동네의 선정되지 않은 "
     "단지들은 거의 움직이지 않았다. 둔산동 비선정 69개 단지는 2.1% 늘었고 법동 8개 단지는 1.5% "
     "증가했으며 탄방동 33개 단지는 2.8% 줄어드는 데 그쳤다. 매물이 줄어든 경로도 통념과 달랐다. "
     "발표 후 엿새간 신규 등록이 61건에서 39건으로 36% 급감한 반면 회수는 69건에서 82건으로 19% "
     "느는 데 그쳐, 이미 나온 물건을 거둬들인 것보다 새 물건이 나오지 않은 영향이 더 컸다. 대전은 "
     "아직 발표 후 실거래 신고가 접수되지 않아(신고 기한 30일) 거래량·가격은 다음 달부터 확인된다."),
    ("해석 시 유의점",
     "이번 비교는 앞선 두 지역이 지나온 초기 3개월의 경로를 보여주는 것이며, 대전의 향후 거래량이나 "
     "가격을 예측하지 않는다. 분당은 수도권 1기 신도시, 부산·대전은 지방권으로 시장 규모와 성격이 "
     "다르고, 세 지역의 관찰 기간에 적용된 금융·세제 환경도 동일하지 않다. 대상 단지들의 월 거래는 "
     "분당 15~82건, 부산 11~38건으로 편차가 커서 특정 시점의 단일 수치보다 방향과 범위로 읽는 것이 "
     "적절하다. 선정 4개월 이후 구간은 기준 기간에 따라 결과가 엇갈려 이번 자료에서는 다루지 않았다."),
    ("일자별 매물·호가, 누구나 단지 단위로 확인 가능",
     "콕집은 전국 아파트·오피스텔 6만 4천여 단지의 매물과 실거래를 매일 수집해 교차 분석하는 "
     "플랫폼으로, 이번 분석에 쓰인 일자별 매물수·실매물수·평균 호가 추이는 콕집 단지 상세의 "
     "‘매물분석’ 메뉴에서 누구나 무료로 확인할 수 있다. 지역 비교, 부동산 타임머신(정책 연대기), "
     "급매 탐지, 단지별 신고가 등 분석 기능도 함께 제공한다."),
]

QUOTE = ("조용호 콕집 공동창업자는 “선도지구로 지정되면 곧바로 값이 뛴다고들 하지만, 앞서 지정된 "
         "두 지역에서 먼저 나타난 것은 거래가 멈추는 현상이었다”며 “분당은 지정 두 달 만에 거래가 "
         "선정 전의 3분의 1 수준까지 줄었고, 부산도 석 달째 감소가 이어졌다. 그 사이 가격은 두 지역 "
         "모두 곧바로 올라 유지됐다. 지금 대전에서 매물이 잠기는 것도 같은 초기 국면으로 보인다”고 "
         "말했다.")

COMPANY = [
    ("서비스", "콕집 — 부동산 매물·실거래·중개사 분석 플랫폼 (koczip.com)"),
    ("운영사", "런투온라인 (대표 황인찬)"),
    ("문의", "조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946"),
]

METHOD = (
    "분석 방법 | 대상: ①분당 3개 구역 12개 단지(샛별마을 동성·라이프·삼부·우방, 양지마을 금호1·청구2·한양5·"
    "금호한양3,5·금호청구6·한양6, 시범단지 우성·현대) 10,738세대, 2024-11-27 선정 ②부산 2개 구역 5개 단지"
    "(코오롱하늘채1·2차, 두산1차, LG, 대림1차) 7,318세대, 2025-12-12 선정 ③대전 3개 구역 6개 단지"
    "(한가람·공작한양·목련·크로바·보람·삼익소월) 7,797세대, 2026-07-15 선정. "
    "실거래 건수·평균가는 국토교통부 실거래 신고 자료(해제거래 제외) 기준이며 신고 기한 30일로 최근 2개월은 미완성. "
    "비교 기준값은 각 지역 선정 직전 3·6·12개월 월평균을 각각 적용해 산출한 결과의 범위. "
    "매물은 콕집 일별 수집 기준(광고매물=포털 노출 광고 건수, 실매물=동일 주택의 중복 광고를 합친 수). "
    "전 수치는 콕집 자체 구축 DB 실측(2026-07-20)."
)

# ── 실측 데이터(2026-07-20 DB) ──
# 기준선을 '선정 직전 3·6·12개월' 세 가지로 각각 잡아 계산한 결과의 범위.
# 이 단지들은 월 거래가 20~30건대라 기준 구간을 어떻게 잡느냐에 따라 수치가 흔들린다.
# 그래서 한 숫자를 단정하지 않고 **세 기준 모두에서 같은 방향인지**를 본다.
#   (K개월, 최소%, 최대%)  — 세 기준(3·6·12개월) 결과의 최소~최대
TRADE_RANGE = {
    "분당": [(1, -59.4, -37.0), (2, -71.3, -55.6), (3, -59.4, -37.0)],
    "부산": [(1, -44.4, -10.4), (2, -54.2, -26.1), (3, -54.6, -26.9)],
}
PRICE_RANGE = {
    "분당": [(1, 8.4, 15.9), (2, 7.2, 14.6), (3, 10.6, 18.2)],
    "부산": [(1, 9.7, 11.4), (2, 8.9, 10.5), (3, 8.7, 10.3)],
}

TABLE = [   # 기준 3·6·12개월 결과의 범위로 표기
    ("선정일", "2024-11-27", "2025-12-12", "2026-07-15"),
    ("규모", "12단지 10,738세대", "5단지 7,318세대", "6단지 7,797세대"),
    ("+1개월 거래량", "-37 ~ -59%", "-10 ~ -44%", "신고 전"),
    ("+2개월 거래량", "-56 ~ -71%", "-26 ~ -54%", "신고 전"),
    ("+3개월 거래량", "-37 ~ -59%", "-27 ~ -55%", "신고 전"),
    ("+1개월 가격", "+8 ~ +16%", "+10 ~ +11%", "신고 전"),
    ("+2개월 가격", "+7 ~ +15%", "+9 ~ +11%", "신고 전"),
    ("+3개월 가격", "+11 ~ +18%", "+9 ~ +10%", "신고 전"),
]


def make_chart():
    """기준선을 3·6·12개월 세 가지로 바꿔 계산한 **결과 범위**를 막대로 보여준다.
    이 단지들은 월 거래가 20~30건대여서 기준 구간에 따라 수치가 흔들린다. 한 숫자를 단정하는
    대신 '어느 기준으로 계산해도 방향이 같다'를 보이는 편이 정확하고 근거도 강하다.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.4, 4.0), dpi=200)
    W = 0.34
    COL = {"분당": "#1268d3", "부산": "#e08a1e"}

    for ax, DATA, title, ylim, legloc in (
        (ax1, TRADE_RANGE, "① 거래량 — 어느 기준으로 계산해도 줄었다", (-84, 10), "upper right"),
        (ax2, PRICE_RANGE, "② 가격 — 어느 기준으로 계산해도 올랐다", (0, 25), "upper left"),
    ):
        ax.axhline(0, color="#333", lw=1.1)
        for gi, (city, rows) in enumerate(DATA.items()):
            xs = [r[0] + (gi - 0.5) * W for r in rows]
            lo = [r[1] for r in rows]; hi = [r[2] for r in rows]
            heights = [h - l for l, h in zip(lo, hi)]
            ax.bar(xs, heights, W, bottom=lo, color=COL[city], alpha=.85,
                   label=f"{city} (기준 3·6·12개월 결과 범위)")
            for x, l, h in zip(xs, lo, hi):
                near = h if abs(h) >= abs(l) else l
                far = l if abs(h) >= abs(l) else h
                ax.annotate(f"{far:+.0f}~{near:+.0f}%", (x, near),
                            ha="center", va="bottom" if near > 0 else "top",
                            xytext=(0, 5 if near > 0 else -5), textcoords="offset points",
                            fontsize=8, fontweight="bold", color=COL[city])
        ax.set_xticks([1, 2, 3])
        ax.set_xticklabels(["선정 +1개월", "+2개월", "+3개월"], fontsize=9.5)
        ax.set_ylabel("선정 전 대비 (%)", fontsize=9)
        ax.set_ylim(*ylim)
        ax.tick_params(axis="y", labelsize=8.5)
        ax.grid(axis="y", color="#eef2f6", lw=.8)
        for sp in ("top", "right"):
            ax.spines[sp].set_visible(False)
        ax.set_title(title, fontsize=10.5, fontweight="bold", pad=8, loc="left")
        ax.legend(fontsize=8, frameon=False, loc=legloc)

    fig.suptitle("선도지구 선정 직후 3개월 — 분당·부산 (콕집 DB)",
                 fontsize=11.5, fontweight="bold", y=0.99)
    fig.text(0.5, 0.005,
             "막대는 기준선을 '선정 직전 3개월·6개월·12개월' 세 가지로 각각 잡아 계산한 결과의 범위. "
             "막대가 0선 한쪽에만 있으면 어떤 기준을 써도 방향이 같다는 뜻.",
             ha="center", fontsize=8, color="#64748b")
    fig.tight_layout(rect=[0, 0.04, 1, 0.95])
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 7월 · 즉시 보도 가능    문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(15); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.2))

    h = d.add_paragraph(); r = h.add_run("■ 3개 도시 한눈에 비교 (콕집 DB)")
    r.font.bold = True; r.font.size = Pt(11.5)
    tbl = d.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ["항목", "분당 (1기 신도시)", "부산", "대전"]):
        c.paragraphs[0].add_run(txt).bold = True
    for row in TABLE:
        for c, txt in zip(tbl.add_row().cells, row):
            c.paragraphs[0].add_run(txt)
    for rowx in tbl.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(8.5)

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    m = d.add_paragraph(); r = m.add_run(METHOD)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=14.5, leading=20, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=10.8, textColor=blue, leading=15, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "method": ParagraphStyle("method", fontName="P", fontSize=7.6, leading=11, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=8, leading=11),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=8, leading=11),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=17 * mm, rightMargin=17 * mm,
                            topMargin=15 * mm, bottomMargin=15 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 176 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    def cell(txt, bold=False):
        return Paragraph(txt, ss["cellb"] if bold else ss["cell"])

    tdata = [[cell(x, True) for x in ["항목", "분당 (1기 신도시)", "부산", "대전"]]]
    for row in TABLE:
        tdata.append([cell(row[0], True), cell(row[1]), cell(row[2]), cell(row[3])])
    tbl = Table(tdata, colWidths=[36 * mm, 50 * mm, 46 * mm, 44 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))

    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 7월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
        chart,
        Spacer(1, 4 * mm),
        Paragraph("■ 3개 도시 한눈에 비교 (콕집 DB)", ss["head"]),
        tbl,
        Spacer(1, 2 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph(METHOD, ss["method"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "선도지구_3도시비교.docx")
    make_pdf(HERE / "선도지구_3도시비교.pdf")
