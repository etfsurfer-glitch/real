# -*- coding: utf-8 -*-
"""보도자료 docx + pdf 생성 — python-docx / reportlab(Pretendard).
Run: python3 design/press/make_press.py  → design/press/콕집_보도자료.docx/.pdf"""
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from press_text import TITLE, SUBTITLE, LEAD, BODY, QUOTE, COMPANY  # noqa: E402

FONTS = HERE.parent / "fonts"
SHOT = HERE / "shots"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)

    p = d.add_paragraph()
    r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 7월 · 즉시 보도 가능    문의: runtoonline@gmail.com · 010-5942-8014")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(16); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)
    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")

    h = d.add_paragraph(); r = h.add_run("■ 실제 구동 화면 (첨부 이미지 — 기사 사용 가능)")
    r.font.bold = True; r.font.size = Pt(11.5)
    # 대표 스크린샷 2장 본문 삽입 (일반 급매 / 중개사 매물점검)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    tbl = d.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, (img, cap) in zip(tbl.rows[0].cells, [
            (SHOT / "02_급매찾기_전국.png", "일반회원 — 급매찾기"),
            (SHOT / "07_중개사_매물점검.png", "중개사회원 — 매물 표시·광고 자가점검")]):
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(img), width=Cm(7.2))
        cp = cell.add_paragraph(cap); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9); cp.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=16.5, leading=23, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=11, textColor=blue, leading=15, spaceAfter=12),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "cap": ParagraphStyle("cap", fontName="P", fontSize=8, textColor=gray, alignment=1),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm)
    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 7월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: runtoonline@gmail.com · 010-5942-8014", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))

    story.append(Paragraph("■ 실제 구동 화면 (기사 사용 가능 — 원본 이미지 별도 첨부)", ss["head"]))
    def shot_cell(fn, cap):
        img = RLImage(str(SHOT / fn), width=54 * mm, height=54 * mm * _ratio(SHOT / fn))
        return [img, Paragraph(cap, ss["cap"])]
    def _ratio(p):
        from PIL import Image as PImage
        w, h = PImage.open(p).size
        return h / w
    row = [shot_cell("02_급매찾기_전국.png", "일반 — 급매찾기"),
           shot_cell("03_우리동네중개사_랭킹.png", "일반 — 우리동네 중개사"),
           shot_cell("07_중개사_매물점검.png", "중개사 — 매물 자가점검")]
    tbl = Table([[c for c in row]], colWidths=[58 * mm] * 3)
    tbl.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                             ("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    story.append(tbl)
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_docx(HERE / "콕집_보도자료.docx")
    make_pdf(HERE / "콕집_보도자료.pdf")
