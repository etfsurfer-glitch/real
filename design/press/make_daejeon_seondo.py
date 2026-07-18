# -*- coding: utf-8 -*-
"""대전 선도지구 발표(2026-07-15) 전후 대상단지 매물·호가 분석 보도자료.
수치는 2026-07-18 complex_daily_agg 실측(단지 6곳: 한가람·공작한양·목련·크로바·보람·삼익소월).
Run: python3 design/press/make_daejeon_seondo.py
 → design/press/대전선도지구_보도자료.docx / .pdf + data/daejeon_seondo_chart.png
발송 없음 — 파일 생성만.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent / "fonts"
DATA = HERE / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "daejeon_seondo_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

# ── 원고 ──
TITLE = "발표 사흘 만에 매물 15%가 사라졌다 — 대전 선도지구 6개 단지, 콕집 DB로 본 발표 전후"
SUBTITLE = "콕집, 일자별 매물 전수 분석 — 광고매물 227→193건, 남은 호가는 역대 신고가마저 추월"

LEAD = (
    "국토교통부와 대전광역시가 7월 15일 노후계획도시 정비 선도지구 3개 구역(둔산지구 13·14구역, 송촌지구 6구역) 선정을 발표한 직후, 대상 단지의 매물이 빠르게 회수되고 남은 매물의 호가는 역대 최고 실거래가를 "
    "넘어서기 시작한 것으로 나타났다. 부동산 데이터 플랫폼 콕집(koczip.com)이 자체 구축한 매물·실거래 "
    "데이터베이스로 대상 6개 단지(한가람·공작한양·목련·크로바·보람·삼익소월, 총 7,797가구)를 일자별로 "
    "전수 추적한 결과다. 발표 전날인 7월 14일 227건이던 6개 단지 매매 광고매물은 7월 18일 193건으로 "
    "사흘 만에 34건(15.0%) 줄었다. 같은 집의 중복 광고를 하나로 합친 ‘실매물’ 기준으로도 144건에서 "
    "127건으로 11.8% 감소했다."
)

BODY = [
    ("대상 단지는 어떤 곳인가 — 전 단지 준공 30년 초과, 둔산은 올해 신고가 경신 중",
     "콕집 DB 기준 6개 단지는 모두 1991~1995년 준공으로 재건축 연한(30년)을 훌쩍 넘긴 대단지다. "
     "한가람(1991년, 1,380가구)·공작한양(1992년, 1,074가구)이 둔산 14구역, 목련(1993년, "
     "1,166가구)·크로바(1992년, 1,632가구)가 13구역, 보람(1995년, 1,735가구)·삼익소월(1993년, "
     "810가구)이 송촌지구 6구역(중리·법동 일대)을 이룬다. 주목할 대목은 발표 이전부터 시장이 움직이고 있었다는 "
     "점이다. 둔산 4개 단지는 올해 상반기에 나란히 역대 최고가를 갈아치웠다. 공작한양이 3월 30일 "
     "6억 1,000만 원, 목련이 4월 17일 15억 5,000만 원, 한가람이 4월 28일 6억 2,000만 원, 크로바가 "
     "5월 19일 23억 5,000만 원으로 모두 2026년에 신고가를 새로 썼다. 선도지구 기대감이 이미 가격에 "
     "실리고 있었고, 발표가 그 흐름에 불을 붙인 셈이다."),
    ("발표 전엔 늘던 매물이 발표 직후 급감 — ‘회수 반전’",
     "매물 감소가 발표 이전부터 이어진 흐름이 아니라는 점이 데이터로 확인된다. 발표 전 한 달간(6월 18일~"
     "7월 14일) 6개 단지 광고매물은 217건에서 227건으로 오히려 늘고 있었다. 증가 추세가 발표 당일(7월 "
     "15일, 220건)을 기점으로 꺾여 16일 206건, 17일 195건, 18일 193건으로 사흘 연속 감소했다. 재건축 "
     "기대감이 커지자 매도자들이 내놓았던 집을 거둬들이는 ‘매물 잠김’이 일자별 데이터에 그대로 찍힌 것이다. "
     "회수 강도는 구역별로 달랐다. 둔산 14구역(한가람·공작한양)이 84건에서 65건으로 22.6% 줄어 가장 "
     "민감했고, 단지별로는 한가람이 31건에서 22건(-29.0%)으로 감소 폭이 가장 컸다. 둔산 13구역은 "
     "-13.0%, 송촌지구 6구역은 -7.6%였다. 올해 신고가를 경신한 둔산 쪽이 더 세게 잠겼고, "
     "신고가가 2023년에 머물러 있는 보람(4억 5,000만 원, 2023년 11월)·삼익소월(2억 9,800만 원, "
     "2023년 9월)은 상대적으로 차분했다 — 기대감의 온도차가 매물 회수 강도에도 그대로 나타났다."),
    ("남은 호가는 신고가를 넘었다",
     "매물이 줄어드는 동안 남은 매물의 가격표는 역대 최고 실거래가 위로 올라섰다. 한가람 103A타입 최저 "
     "호가는 발표 전 6억 5,000만 원에서 6억 8,000만 원으로 올라, 이 단지 역대 신고가(6억 2,000만 원)를 "
     "9.7% 웃돈다. 크로바는 189타입 최저 호가가 24억 원에서 24억 5,000만 원으로 상향돼 역대 신고가 "
     "(23억 5,000만 원)보다 1억 원 높아졌고, 단지 전체 평균 호가(매물수 가중)도 17억 5,200만 원에서 "
     "17억 8,300만 원으로 사흘 새 1.8% 올랐다 — 최근 6개월 실거래 평균(16억 3,600만 원)보다 9.0% 높은 "
     "수준이다. 이 밖에 보람 117B타입 최저 호가가 2억 9,000만 원에서 3억 2,000만 원(+10.3%), 목련 "
     "122A타입이 11억 5,000만 원에서 12억 원(+4.3%)으로 상향됐다."),
    ("전세 매물은 그대로 — ‘팔 사람’만 멈췄다",
     "같은 기간 6개 단지의 전세 매물은 46건에서 50건으로 줄지 않았다. 잠긴 것은 매매뿐이다. 거주 수요가 "
     "빠져나간 것이 아니라, 소유자들이 값이 더 오를 것으로 보고 매도만 보류하고 있다는 뜻으로 읽히는 "
     "대목이다."),
    ("해석 시 유의점",
     "이번 수치는 발표 후 사흘간의 초기 반응이다. 매물 감소에는 호가를 다시 부르기 위한 회수와 실제 계약 "
     "성사가 섞여 있을 수 있다. 발표 이후 계약분의 실거래 신고는 콕집 DB에 아직 접수되지 않았는데(신고 "
     "기한 30일), 신고가를 넘어선 호가가 실제 거래로 이어질지는 이 자료가 쌓여야 확인된다. 다만 광고 "
     "건수가 아니라 중복 광고를 합친 실매물 기준으로도 같은 방향의 감소가 확인돼, 착시가 아닌 실제 매물 "
     "잠김이라는 점은 분명하다."),
    ("일자별 매물·호가, 누구나 단지 단위로 확인 가능",
     "콕집은 전국 아파트·오피스텔 6만 4천여 단지의 매물과 실거래를 매일 수집해 교차 분석하는 플랫폼으로, "
     "이번 분석에 쓰인 일자별 매물수·실매물수·평균 호가 추이는 콕집 단지 상세의 ‘매물분석’ 메뉴에서 "
     "누구나 무료로 확인할 수 있다. 급매 탐지, 단지별 신고가, 우리동네 중개사 랭킹 등 분석 기능도 함께 "
     "제공한다."),
]

QUOTE = ("조용호 콕집 공동창업자는 “정비사업 발표가 시장에 어떤 파장을 만드는지는 지금까지 체감과 소문으로만 "
         "이야기돼 왔다. 일자별 매물 데이터는 그 반응을 숫자로 보여준다. 소비자가 기대감에 휩쓸리기 전에 "
         "매물과 가격의 실제 움직임을 확인하고 판단하도록 돕겠다”고 말했다.")

COMPANY = [
    ("서비스", "콕집 — 부동산 매물·실거래·중개사 분석 플랫폼 (koczip.com)"),
    ("운영사", "런투온라인 (대표 황인찬)"),
    ("문의", "조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946"),
]

METHOD = (
    "분석 방법 | 대상: 대전 선도지구 3개 구역(둔산 13·14, 송촌 6) 6개 단지(한가람·공작한양·목련·크로바·보람·삼익소월) 매매·전세 매물 및 "
    "실거래. 기간: 2026-06-17~07-18 일자별. 광고매물=포털 노출 광고 건수, 실매물=동일 주택의 중복 광고를 합친 수, "
    "평균 호가=매물수 가중평균. 신고가·실거래 평균은 국토교통부 실거래 신고 자료(해제거래 제외) 기준. 호가는 매도 "
    "희망가로 실거래 가격이 아님. 준공연도·세대수 포함 전 수치는 콕집 자체 구축 DB 실측(2026-07-18)."
)

# ── 실측 데이터(2026-07-18 DB) ──
DAILY = [  # (날짜, 광고매물, 실매물) — 6개 단지 합계. 6/28 실매물은 스냅샷 결측.
    ("06-17", 220, 143), ("06-18", 217, 143), ("06-19", 221, 143), ("06-20", 225, 144),
    ("06-21", 218, 140), ("06-22", 228, 137), ("06-23", 226, 141), ("06-24", 227, 142),
    ("06-25", 232, 142), ("06-26", 228, 146), ("06-27", 228, 144), ("06-28", 224, None),
    ("06-29", 227, 139), ("06-30", 221, 139), ("07-01", 226, 141), ("07-02", 223, 143),
    ("07-03", 228, 140), ("07-04", 232, 144), ("07-05", 229, 144), ("07-06", 235, 142),
    ("07-07", 238, 147), ("07-08", 237, 148), ("07-09", 235, 146), ("07-10", 231, 141),
    ("07-11", 232, 144), ("07-12", 223, 140), ("07-13", 231, 147), ("07-14", 227, 144),
    ("07-15", 220, 137), ("07-16", 206, 134), ("07-17", 195, 128), ("07-18", 193, 127),
]

TABLE = [  # 구역, 단지(위치·준공), 광고 7/14→7/18, 증감, 실매물 7/14→7/18 — 전부 콕집 DB 실측
    ("둔산 13", "목련 (둔산동 · 1993)", "23 → 17", "-26.1%", "15 → 12"),
    ("둔산 13", "크로바 (둔산동 · 1992)", "54 → 50", "-7.4%", "28 → 26"),
    ("둔산 14", "한가람 (탄방동 · 1991)", "31 → 22", "-29.0%", "25 → 19"),
    ("둔산 14", "공작한양 (탄방동 · 1992)", "53 → 43", "-18.9%", "31 → 29"),
    ("송촌 6", "보람 (법동 · 1995)", "30 → 26", "-13.3%", "20 → 17"),
    ("송촌 6", "삼익소월 (법동 · 1993)", "36 → 35", "-2.8%", "25 → 24"),
    ("합계", "6개 단지 (7,797가구)", "227 → 193", "-15.0%", "144 → 127"),
]


def make_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    xs = list(range(len(DAILY)))
    n = [d[1] for d in DAILY]
    u = [d[2] for d in DAILY]
    fig, ax = plt.subplots(figsize=(8.6, 3.4), dpi=200)
    ann = next(i for i, d in enumerate(DAILY) if d[0] == "07-15")
    ax.axvline(ann, color="#e2574c", lw=1.2, ls="--", alpha=.9)
    ax.text(ann - 0.4, 244, "7/15 선도지구 발표", color="#e2574c", fontsize=9,
            fontweight="bold", ha="right", va="top")
    ax.plot(xs, n, color="#1268d3", lw=2.2, label="광고매물")
    # 실매물 결측(6/28)은 선을 끊지 않게 구간 분리 없이 마스크 처리
    ux = [x for x, v in zip(xs, u) if v is not None]
    uv = [v for v in u if v is not None]
    ax.plot(ux, uv, color="#8aa7c9", lw=1.8, label="실매물(중복 합침)")
    ax.fill_between(xs[ann:], n[ann:], color="#e2574c", alpha=.07)
    ax.annotate(f"{n[-1]}건", (xs[-1], n[-1]), textcoords="offset points", xytext=(4, 4),
                fontsize=9.5, fontweight="bold", color="#1268d3")
    ax.annotate(f"{uv[-1]}건", (ux[-1], uv[-1]), textcoords="offset points", xytext=(4, -12),
                fontsize=9, color="#5a7396")
    ticks = [i for i, d in enumerate(DAILY) if d[0].endswith(("17", "24", "01", "08", "14", "15", "18")) and i % 1 == 0]
    ticks = [i for i in ticks if i in (0, 7, 14, 21, 28, 31)]
    ax.set_xticks(ticks)
    ax.set_xticklabels([DAILY[i][0].replace("0", "", 1).replace("-", ".") for i in ticks], fontsize=8.5)
    ax.tick_params(axis="y", labelsize=8.5)
    ax.set_ylim(110, 260)
    ax.grid(axis="y", color="#e8edf3", lw=.7)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    ax.legend(fontsize=9, frameon=False, loc="lower left")
    ax.set_title("대전 선도지구 6개 단지 매매 매물 추이 (2026.6.17~7.18, 콕집 DB)", fontsize=10.5, fontweight="bold", pad=8)
    fig.tight_layout()
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 7월 · 즉시 보도 가능    문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(15.5); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.2))

    h = d.add_paragraph(); r = h.add_run("■ 단지별 개요·매매 매물 변화 (발표 전일 7/14 → 7/18, 콕집 DB)")
    r.font.bold = True; r.font.size = Pt(11.5)
    tbl = d.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ["구역", "단지", "광고매물", "증감", "실매물"]):
        c.paragraphs[0].add_run(txt).bold = True
    for row in TABLE:
        cells = tbl.add_row().cells
        for c, txt in zip(cells, row):
            run = c.paragraphs[0].add_run(txt)
            if row[0] == "합계":
                run.bold = True
    for rowx in tbl.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9.5)

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    m = d.add_paragraph(); r = m.add_run(METHOD)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=15.5, leading=22, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=10.8, textColor=blue, leading=15, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "method": ParagraphStyle("method", fontName="P", fontSize=7.8, leading=11.5, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=8.8, leading=12),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=8.8, leading=12),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 174 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    def cell(txt, bold=False):
        return Paragraph(txt, ss["cellb"] if bold else ss["cell"])

    tdata = [[cell(x, True) for x in ["구역", "단지", "광고매물", "증감", "실매물"]]]
    for row in TABLE:
        tdata.append([cell(x, row[0] == "합계") for x in row])
    tbl = Table(tdata, colWidths=[34 * mm, 44 * mm, 34 * mm, 24 * mm, 34 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("BACKGROUND", (0, len(tdata) - 1), (-1, len(tdata) - 1), colors.Color(0.96, 0.97, 0.985)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))

    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 7월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
        chart,
        Spacer(1, 4 * mm),
        Paragraph("■ 단지별 개요·매매 매물 변화 (발표 전일 7/14 → 7/18, 콕집 DB)", ss["head"]),
        tbl,
        Spacer(1, 2 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph(METHOD, ss["method"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "대전선도지구_보도자료.docx")
    make_pdf(HERE / "대전선도지구_보도자료.pdf")
