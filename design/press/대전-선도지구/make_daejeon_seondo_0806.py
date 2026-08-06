# -*- coding: utf-8 -*-
"""대전 선도지구 발표(2026-07-15) 이후 3주 — 대상단지 매물·호가 현행화 보도자료.

7/24자 배포본의 후속. 그때는 '매물 40% 감소 · 호가는 보합'이었는데, 3주가 지나며
①매물 잠김이 7/31 바닥(113건)을 찍고 멈췄고 ②호가가 실제로 올랐다(구성효과를 걷어낸
고정가중 기준 +7.3%). 프레임이 바뀌는 지점이라 새 자료로 낸다.
수치는 2026-08-06 complex_daily_agg·transactions 실측(단지 6곳).
Run: python3 design/press/대전-선도지구/make_daejeon_seondo_0806.py
 → design/press/대전-선도지구/대전선도지구_보도자료_20260806.docx / .pdf
발송 없음 — 파일 생성만.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent.parent / "fonts"
DATA = HERE.parent / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "daejeon_seondo_chart_0806.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

# ── 원고 ──
TITLE = "매물 잠김은 멈췄는데 호가가 올랐다 — 대전 선도지구 6개 단지, 발표 3주 뒤"
SUBTITLE = "콕집 일자별 전수 분석 — 광고매물 227→121건(-46.7%), 7/31 바닥 뒤 보합 · 호가는 구성효과 걷어내도 +7.3%"

LEAD = (
    "국토교통부와 대전광역시가 7월 15일 노후계획도시 정비 선도지구 3개 구역(둔산지구 13·14구역, 송촌지구 "
    "6구역)을 발표한 지 3주가 지났다. 부동산 데이터 플랫폼 콕집(koczip.com)이 대상 6개 단지"
    "(한가람·공작한양·목련·크로바·보람·삼익소월, 총 7,797가구)를 일자별로 전수 추적한 결과, 발표 직후 "
    "쏟아지던 매물 회수는 7월 31일 113건에서 멈췄고 이후 121건까지 소폭 되돌아왔다. 발표 전날인 7월 14일 "
    "227건과 견주면 여전히 46.7% 적은 수준이다. 같은 집의 중복 광고를 하나로 합친 ‘실매물’ 기준으로도 "
    "144건에서 86건으로 40.3% 줄었다. 달라진 것은 가격이다. 7월 24일자 분석에서 보합이던 호가가 "
    "3주 사이 뚜렷하게 올랐다. 매물 구성이 바뀐 효과를 걷어낸 고정가중 기준으로도 8억 3,100만 원에서 "
    "8억 9,100만 원으로 7.3% 상승했다."
)

BODY = [
    ("매물 회수는 7월 31일에 멈췄다 — 잠김의 1차 국면 종료",
     "발표 당일 220건이던 6개 단지 광고매물은 열이레 연속 줄어 7월 31일 113건까지 내려갔다. 이후 8월 들어 "
     "114~121건 사이에서 오르내리며 더 줄지 않고 있다. 8월 6일 121건은 저점 대비 8건 늘어난 수치다. "
     "발표 직후의 일괄 회수가 일단락되고, 호가를 올려 다시 내놓는 매물이 조금씩 나오는 국면으로 읽힌다. "
     "실매물도 7월 31일 80건에서 8월 6일 86건으로 같은 방향을 그렸다. 다만 발표 전 수준(144건)에는 크게 "
     "못 미쳐, 잠긴 매물이 풀린 것이 아니라 감소가 멈춘 단계로 보는 편이 정확하다."),
    ("호가는 올랐다 — 구성효과를 걷어내도 7.3%",
     "6개 단지 통합 평균 매매 호가는 7월 14일 8억 3,100만 원에서 8월 6일 9억 5,700만 원으로 15.2% 올랐다. "
     "다만 이 값에는 저가 매물이 먼저 회수돼 남은 매물이 비싸 보이는 구성효과가 섞여 있다. 이를 걷어내기 위해 "
     "발표 전(7/14) 단지별 매물수를 가중치로 고정해 다시 계산하면 8억 3,100만 원에서 8억 9,100만 원으로 "
     "7.3% 상승한다. 절반 정도는 구성효과이고, 나머지 절반은 실제 호가 상승이라는 뜻이다. 단지별로도 6곳 중 "
     "5곳이 올랐다. 목련이 13억 5,100만 원에서 16억 200만 원으로 18.5% 올라 상승폭이 가장 컸고, 공작한양 "
     "7.0%, 크로바 5.5%, 삼익소월 4.6%, 한가람 4.5% 순이었다. 보람만 3억 3,600만 원에서 3억 2,100만 "
     "원으로 4.6% 내렸다. 7월 24일자 분석에서 ‘아직 보합’이던 호가가 3주 사이 방향을 잡은 셈이다."),
    ("구역별로 갈렸다 — 둔산 14구역이 가장 세게 잠겼다",
     "회수 강도는 구역별로 크게 달랐다. 둔산 14구역(한가람·공작한양)이 84건에서 27건으로 67.9% 줄어 가장 "
     "강하게 잠겼고, 둔산 13구역(목련·크로바)은 77건에서 50건으로 35.1%, 송촌 6구역(보람·삼익소월)은 "
     "66건에서 44건으로 33.3% 감소했다. 단지 단위로는 공작한양이 53건에서 13건으로 75.5% 줄어 감소폭이 "
     "가장 컸다. 공작한양은 발표 직후인 7월 20일 6억 7,000만 원에 역대 최고가를 새로 쓴 단지다. "
     "가격이 확인된 곳일수록 매도자가 더 오래 기다리는 모습이 데이터에 그대로 나타난다."),
    ("발표 후 실거래 13건 신고 — 신고가는 아직 깨지지 않았다",
     "7월 15일 이후 계약분으로 신고된 실거래는 6개 단지 합쳐 13건이다. 공작한양 3건(최고 6억 7,000만 원, "
     "7/20), 한가람 3건(최고 6억 원, 7/28), 삼익소월 3건(최고 2억 1,800만 원, 7/31), 크로바 2건(최고 "
     "20억 원, 7/22), 보람 2건(최고 3억 2,000만 원, 7/30)이다. 목련은 아직 신고분이 없다. 크로바의 20억 "
     "원은 5월에 세운 역대 최고가 23억 5,000만 원에 못 미치고, 목련(15억 5,000만 원, 4월)·한가람"
     "(6억 2,000만 원, 4월)의 최고가도 아직 깨지지 않았다. 호가는 올랐지만 그 값에 실제 계약이 붙는지는 "
     "신고 기한(30일)이 지나 8월 하순 이후 신고분이 쌓여야 확인된다."),
    ("전세는 계속 늘었다 — 잠긴 것은 여전히 매매뿐",
     "같은 기간 6개 단지의 전세 매물은 46건에서 66건으로 43.5% 늘었다. 7월 24일 54건에서 다시 12건이 "
     "더해진 것이다. 실매물 기준으로도 28건에서 50건으로 늘었다. 매매만 잠기고 임대 물량은 오히려 풀리는 "
     "흐름이 3주째 이어지고 있다. 소유자들이 집을 파는 대신 세를 놓으며 사업 진행을 기다리는 선택으로 "
     "읽히는 대목이다."),
    ("해석 시 유의점",
     "이번 수치는 발표 후 3주(8월 6일 기준)의 반응이다. 매물 감소에는 호가를 다시 부르기 위한 회수와 실제 "
     "계약 성사가 섞여 있다. 호가 상승분 역시 매물이 크게 줄어든 상태에서 산출된 값이라 표본이 얇다는 점을 "
     "감안해야 한다. 고정가중 계산은 단지 구성 변화만 보정할 뿐, 같은 단지 안에서 대형 평형만 남는 효과까지 "
     "제거하지는 못한다. 실거래로 확인되기 전까지 호가는 매도 희망가일 뿐이다."),
    ("일자별 매물·호가, 누구나 단지 단위로 확인 가능",
     "콕집은 전국 아파트·오피스텔 6만 4천여 단지의 매물과 실거래를 매일 수집해 교차 분석하는 플랫폼으로, "
     "이번 분석에 쓰인 일자별 매물수·실매물수·평균 호가 추이는 콕집 단지 상세의 ‘매물분석’ 메뉴에서 "
     "누구나 무료로 확인할 수 있다. 급매 탐지, 단지별 신고가, 우리동네 중개사 랭킹 등 분석 기능도 함께 "
     "제공한다."),
]

QUOTE = ("조용호 콕집 공동창업자는 “발표 직후 3주는 매물이 잠기는 국면이었고, 지금은 남은 매물의 값이 오르는 "
         "국면으로 넘어가고 있다. 통합 평균만 보면 15% 올랐지만 절반은 매물 구성이 바뀐 착시다. 일자별 "
         "데이터로 그 둘을 갈라 보여주는 것이 소비자가 판단할 때 필요한 정보라고 본다”고 말했다.")

COMPANY = [
    ("서비스", "콕집 — 부동산 매물·실거래·중개사 분석 플랫폼 (koczip.com)"),
    ("운영사", "런투온라인 (대표 황인찬)"),
    ("문의", "조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946"),
]

METHOD = (
    "분석 방법 | 대상: 대전 선도지구 3개 구역(둔산 13·14, 송촌 6) 6개 단지(한가람·공작한양·목련·크로바·"
    "보람·삼익소월) 매매·전세 매물 및 실거래. 기간: 2026-06-17~08-06 일자별. 광고매물=포털 노출 광고 건수, "
    "실매물=동일 주택의 중복 광고를 합친 수, 평균 호가=매물수 가중평균. 고정가중 호가=발표 전일(7/14) 단지별 "
    "매물수를 가중치로 고정해 단지 구성 변화 효과를 제거한 값. 신고가·실거래는 국토교통부 실거래 신고 자료"
    "(해제거래 제외) 기준. 호가는 매도 희망가로 실거래 가격이 아님. 전 수치는 콕집 자체 구축 DB 실측"
    "(2026-08-06)."
)

# ── 실측 데이터(2026-08-06 DB) ──
DAILY = [
    ('06-17', 220, 143),
    ('06-18', 217, 143),
    ('06-19', 221, 143),
    ('06-20', 225, 144),
    ('06-21', 218, 140),
    ('06-22', 228, 137),
    ('06-23', 226, 141),
    ('06-24', 227, 142),
    ('06-25', 232, 142),
    ('06-26', 228, 146),
    ('06-27', 228, 144),
    ('06-28', 224, None),
    ('06-29', 227, 139),
    ('06-30', 221, 139),
    ('07-01', 226, 141),
    ('07-02', 223, 143),
    ('07-03', 228, 140),
    ('07-04', 232, 144),
    ('07-05', 229, 144),
    ('07-06', 235, 142),
    ('07-07', 238, 147),
    ('07-08', 237, 148),
    ('07-09', 235, 146),
    ('07-10', 231, 141),
    ('07-11', 232, 144),
    ('07-12', 223, 140),
    ('07-13', 231, 147),
    ('07-14', 227, 144),
    ('07-15', 220, 137),
    ('07-16', 206, 134),
    ('07-17', 195, 128),
    ('07-18', 191, 126),
    ('07-19', 187, 126),
    ('07-20', 176, 117),
    ('07-21', 169, 114),
    ('07-22', 158, 110),
    ('07-23', 142, 101),
    ('07-24', 126, 87),
    ('07-25', 127, 88),
    ('07-26', 124, 85),
    ('07-27', 122, 82),
    ('07-28', 124, 87),
    ('07-29', 118, 83),
    ('07-30', 117, 83),
    ('07-31', 113, 80),
    ('08-01', 115, 82),
    ('08-02', 114, 81),
    ('08-03', 114, 81),
    ('08-04', 116, 84),
    ('08-05', 121, 87),
    ('08-06', 121, 86),
]

TABLE = [  # 구역, 단지(위치·준공), 광고 7/14→8/06, 증감, 호가 7/14→8/06(억) — 전부 콕집 DB 실측
    ("둔산 13", "목련 (둔산동 · 1993)", "23 → 13", "-43.5%", "13.51 → 16.02"),
    ("둔산 13", "크로바 (둔산동 · 1992)", "54 → 37", "-31.5%", "17.52 → 18.49"),
    ("둔산 14", "한가람 (탄방동 · 1991)", "31 → 14", "-54.8%", "4.60 → 4.80"),
    ("둔산 14", "공작한양 (탄방동 · 1992)", "53 → 13", "-75.5%", "6.03 → 6.45"),
    ("송촌 6", "보람 (법동 · 1995)", "30 → 23", "-23.3%", "3.36 → 3.21"),
    ("송촌 6", "삼익소월 (법동 · 1993)", "36 → 21", "-41.7%", "1.86 → 1.95"),
    ("합계", "6개 단지 (7,797가구)", "227 → 121", "-46.7%", "8.31 → 9.57"),
]


def make_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    xs = list(range(len(DAILY)))
    n = [d[1] for d in DAILY]
    u = [d[2] for d in DAILY]
    fig, ax = plt.subplots(figsize=(8.6, 3.4), dpi=200)
    ann = next(i for i, d in enumerate(DAILY) if d[0] == "07-15")
    ax.axvline(ann, color="#e2574c", lw=1.2, ls="--", alpha=.9)
    ax.text(ann - 0.4, 244, "7/15 선도지구 발표", color="#e2574c", fontsize=9,
            fontweight="bold", ha="right", va="top")
    ax.plot(xs, n, color="#1268d3", lw=2.2, label="광고매물")
    # 실매물 결측(6/28)은 선을 끊지 않게 구간 분리 없이 마스크 처리
    ux = [x for x, v in zip(xs, u) if v is not None]
    uv = [v for v in u if v is not None]
    ax.plot(ux, uv, color="#8aa7c9", lw=1.8, label="실매물(중복 합침)")
    ax.fill_between(xs[ann:], n[ann:], color="#e2574c", alpha=.07)
    ax.annotate(f"{n[-1]}건", (xs[-1], n[-1]), textcoords="offset points", xytext=(4, 4),
                fontsize=9.5, fontweight="bold", color="#1268d3")
    ax.annotate(f"{uv[-1]}건", (ux[-1], uv[-1]), textcoords="offset points", xytext=(4, -12),
                fontsize=9, color="#5a7396")
    ticks = [i for i, d in enumerate(DAILY) if d[0] in ("06-17", "06-24", "07-01", "07-08", "07-15", "07-22", "07-29", "08-06")]
    ax.set_xticks(ticks)
    ax.set_xticklabels([DAILY[i][0].replace("0", "", 1).replace("-", ".") for i in ticks], fontsize=8.5)
    ax.tick_params(axis="y", labelsize=8.5)
    ax.set_ylim(70, 260)
    ax.grid(axis="y", color="#e8edf3", lw=.7)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    ax.legend(fontsize=9, frameon=False, loc="lower left")
    ax.set_title("대전 선도지구 6개 단지 매매 매물 추이 (2026.6.17~8.6, 콕집 DB)", fontsize=10.5, fontweight="bold", pad=8)
    fig.tight_layout()
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 8월 · 즉시 보도 가능    문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(15.5); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.2))

    h = d.add_paragraph(); r = h.add_run("■ 단지별 매매 매물·호가 변화 (발표 전일 7/14 → 8/06, 콕집 DB)")
    r.font.bold = True; r.font.size = Pt(11.5)
    tbl = d.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ["구역", "단지", "광고매물", "증감", "평균 호가(억)"]):
        c.paragraphs[0].add_run(txt).bold = True
    for row in TABLE:
        cells = tbl.add_row().cells
        for c, txt in zip(cells, row):
            run = c.paragraphs[0].add_run(txt)
            if row[0] == "합계":
                run.bold = True
    for rowx in tbl.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9.5)

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    m = d.add_paragraph(); r = m.add_run(METHOD)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=15.5, leading=22, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=10.8, textColor=blue, leading=15, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "method": ParagraphStyle("method", fontName="P", fontSize=7.8, leading=11.5, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=8.8, leading=12),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=8.8, leading=12),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 174 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    def cell(txt, bold=False):
        return Paragraph(txt, ss["cellb"] if bold else ss["cell"])

    tdata = [[cell(x, True) for x in ["구역", "단지", "광고매물", "증감", "평균 호가(억)"]]]
    for row in TABLE:
        tdata.append([cell(x, row[0] == "합계") for x in row])
    tbl = Table(tdata, colWidths=[34 * mm, 44 * mm, 34 * mm, 24 * mm, 34 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("BACKGROUND", (0, len(tdata) - 1), (-1, len(tdata) - 1), colors.Color(0.96, 0.97, 0.985)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))

    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 8월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
        chart,
        Spacer(1, 4 * mm),
        Paragraph("■ 단지별 매매 매물·호가 변화 (발표 전일 7/14 → 8/06, 콕집 DB)", ss["head"]),
        tbl,
        Spacer(1, 2 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph(METHOD, ss["method"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "대전선도지구_보도자료_20260806.docx")
    make_pdf(HERE / "대전선도지구_보도자료_20260806.pdf")
