# -*- coding: utf-8 -*-
"""삼성 탕정 반도체벨트 배후 4개 주거지 대장아파트 — 84㎡ 실거래 추이 (취재 참고자료).
서울경제 회신 대응: 반도체벨트 인구·집값 기사의 받침 박스용. 대장단지 위주 거래량·신고가·가격.
수치는 2026-07-28 transactions·complex_daily_agg 실측(반기별). 신고 30일 지연으로 최근 구간 미완성.
Run: python3 design/press/반도체벨트/make_semibelt_flagship.py
 → design/press/반도체벨트/반도체벨트_대장아파트.docx / .pdf + data/semibelt_flagship_chart.png

※ 보도자료가 아니라 취재 참고자료(보고서). 예측·단정 금지, 관측된 사실과 한계만.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent.parent / "fonts"
DATA = HERE.parent / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "semibelt_flagship_chart.png"
CHART2 = DATA / "semibelt_maemul_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)
COLORS = {"불당 지웰더샵": "#1268d3", "탕정 지웰시티": "#e08a1e",
          "성성 레이크사이드": "#1f9d63", "배방 e편한세상탕정": "#d23b3b"}

DOCTYPE = "데이터 분석 보고서 (취재 참고자료)"
TITLE = "삼성 탕정 반도체벨트 배후 4개 주거지 대장아파트 — 성숙 단지는 신고가·거래 견조, 신축 입주장은 매물 소화"
SUBTITLE = ("천안 불당동·성성동, 아산 탕정지구·배방지구 — 각 지역 대장아파트 84㎡(국민평형) 실거래 추이 · "
            "2026년 7월 28일 기준 · 실거래 신고 30일 지연으로 2026년 상반기 이후 구간은 미완성")

# 반기 축 (2023하 ~ 2026상, 6구간)
PERIODS = ["2023하", "2024상", "2024하", "2025상", "2025하", "2026상"]

# 대장단지별 시계열 (84㎡ 전용, 억원 / 반기 전체거래 건)
SERIES = {
    "불당 지웰더샵": {
        "hi":  [8.30, 8.47, 8.47, 8.50, 8.70, 9.46],
        "avg": [8.07, 8.30, 8.21, 8.20, 8.38, 8.61],
        "vol": [13, 19, 27, 16, 25, 28],
        "color": "#1268d3",
    },
    "탕정 지웰시티": {
        "hi":  [7.70, 7.50, 7.40, 7.30, 7.40, 7.48],
        "avg": [6.91, 6.87, 6.86, 6.91, 6.92, 6.93],
        "vol": [28, 62, 64, 56, 94, 103],
        "color": "#e08a1e",
    },
    "성성 레이크사이드": {
        "hi":  [5.30, 5.90, 6.15, 6.30, 6.25, 5.89],
        "avg": [5.30, 5.85, 5.87, 5.95, 5.85, 5.67],
        "vol": [2, 3, 3, 15, 41, 34],
        "color": "#1f9d63",
    },
    "배방 e편한세상탕정": {  # 2025 준공 — 2025하부터
        "hi":  [None, None, None, None, 4.60, 5.50],
        "avg": [None, None, None, None, 4.40, 5.30],
        "vol": [0, 0, 0, 0, 2, 8],
        "color": "#d23b3b",
    },
}

# ── 매물·호가 (콕집 일별 수집 강점 — 실거래 30일 지연이 못 보여주는 즉각 반응) ──
# 주간(7일) 매매 광고매물 추이 — 2026-06-02 ~ 07-28 (9주)
MAEMUL_WEEKS = ["6/2", "6/9", "6/16", "6/23", "6/30", "7/7", "7/14", "7/21", "7/28"]
MAEMUL_SERIES = {
    "불당 지웰더샵":     [91, 92, 94, 119, 113, 99, 102, 97, 96],
    "탕정 지웰시티":     [701, 687, 646, 627, 600, 482, 353, 267, 241],
    "성성 레이크사이드": [544, 561, 564, 563, 582, 580, 552, 549, 550],
    "배방 e편한세상탕정": [42, 35, 33, 41, 26, 25, 19, 17, 17],
}
# 현재 스냅샷(2026-07-28): (광고매물, 실매물, 광고배율, 9주변화%, 84㎡평균호가억, 84㎡호가매물, 호가-실거래갭%)
MAEMUL_NOW = {
    "불당 지웰더샵":     (96, 38, 2.5, 5, 9.21, 16, -2.6),
    "탕정 지웰시티":     (241, 76, 3.2, -66, 7.70, 1, 0.0),
    "성성 레이크사이드": (550, 144, 3.8, 1, 6.27, 7, -0.5),
    "배방 e편한세상탕정": (17, 7, 2.4, -60, 5.90, 3, 7.3),
}

SUMMARY = [
    ("분석 대상", "삼성 탕정캠퍼스 배후 4개 주거지의 대장아파트(주상복합 제외, 지역 최고가 선도) 한 곳씩 — "
                "불당 지웰더샵(685세대)·탕정 지웰시티 1·2·3단지(2,206세대)·성성 레이크사이드(1,023세대)·배방 e편한세상탕정(893세대)"),
    ("최고가", "불당 지웰더샵 84㎡가 2026년 4월 9.46억으로 역대 신고가를 경신 — 4개 지역 대장 중 절대가격 최고"),
    ("거래 급증", "삼성 캠퍼스 옆 탕정 지웰시티는 반기 거래가 28건→103건으로 늘며 평균 6.9억선 유지(신고가는 7.3~7.7억 박스권)"),
    ("★ 매물 급감", "탕정 지웰시티 매매 광고매물이 9주간 701건→241건(–66%)으로 급감 — 실거래 30일 지연으로 안 보이는 매도 심리를 매물이 먼저 드러냄(콕집 일별 수집)"),
    ("★ 매물 정체 vs 감소", "같은 신축이라도 성성 레이크사이드는 매물 550건대에서 소진되지 않고 정체, 배방 e편한세상탕정은 –60%로 정반대 흐름"),
    ("★ 호가–실거래", "성숙 단지(불당·탕정·성성)는 84㎡ 호가가 실거래에 밀착(–2.6~0%)했고, 신축 배방만 호가가 실거래보다 +7.3% 높아 매도자 눈높이가 앞섬"),
    ("유의 사항", "84㎡=국민평형(34평) 대표비교 · 실거래 신고 30일 지연으로 최근 구간 미완성 · 매물수는 광고 건수(중복 포함) · 신축 두 곳은 표본 얇음 · 대장단지라 지역 전체 평균과 다를 수 있음"),
]

BODY = [
    ("1. 무엇을 봤나",
     "삼성 탕정 반도체·디스플레이 캠퍼스 배후의 4개 주거지 — 천안 서북구 불당동·성성동, 아산 탕정지구(탕정면)·"
     "배방지구(배방읍) — 에서 각 지역의 대장아파트를 한 곳씩 골라 전용 84㎡(국민평형) 실거래를 반기별로 추적했습니다. "
     "대장아파트는 주상복합을 제외하고 그 지역에서 가장 비싼, 시세를 선도하는 아파트를 말합니다. 84㎡로 기준을 통일한 것은 "
     "단지·시기별 평형 구성 차이를 걷어내고 같은 잣대로 비교하기 위해서입니다. 본 자료는 관측된 실거래 사실만 정리하며 "
     "앞으로의 가격을 예측하지 않습니다."),
    ("2. 불당동 — 벨트 최고가, 대장이 신고가를 경신했다",
     "천안 도심(서북구)의 대장인 천안불당지웰더샵(685세대, 2016년 준공)의 84㎡ 신고가는 2023년 하반기 8.30억에서 "
     "2026년 4월 9.46억으로 역대 최고가를 새로 썼습니다. 반기 거래도 13건→28건으로 꾸준하고, 84㎡ 평균가는 "
     "8.07억→8.61억으로 완만히 올랐습니다. 매물 쪽도 안정적입니다 — 매매 광고매물이 최근 9주간 90~120건 사이에서 "
     "움직이며 현재 96건이고, 84㎡ 매물 평균 호가(9.21억)는 실거래 신고가(9.46억)에 거의 붙어 있습니다(–2.6%). "
     "4개 지역 대장 중 절대가격이 가장 높고, 호가와 실거래가 정합적인 성숙 단지의 모습입니다."),
    ("3. 탕정지구 — 삼성 캠퍼스 옆, 거래는 급증하고 매물은 9주째 잠기고 있다",
     "삼성 탕정캠퍼스에 바로 인접한 한들물빛도시 지웰시티센트럴푸르지오(1·2·3단지 합계 2,206세대, 2022년 준공)는 "
     "이번 분석에서 가장 뚜렷한 신호를 보였습니다. 84㎡ 신고가는 입주 초기 2023년 7.70억을 찍은 뒤 7.30~7.48억 "
     "박스권이지만, 반기 거래량은 28건(2023년 하반기)→103건(2026년 상반기)으로 크게 늘었고 평균가는 6.9억선을 "
     "지켰습니다. 더 주목되는 건 매물입니다 — 매매 광고매물이 6월 초 701건에서 7월 28일 241건으로 9주 연속 줄어 "
     "66% 급감했습니다. 실거래(신고 30일 지연)로는 아직 안 보이는 매도자들의 관망·회수가 매물 데이터에 먼저 "
     "드러난 것으로, 거래는 활발한데 내놓은 물건은 빠르게 걷혀 84㎡ 매물은 현재 1건에 불과합니다."),
    ("4. 성성동 — 신축 입주장, 거래는 폭발했으나 매물이 550건대에서 안 빠진다",
     "2023년 준공한 천안푸르지오레이크사이드(1,023세대)는 입주장을 맞아 반기 거래가 3건(2024년 하반기)→41건"
     "(2025년 하반기)으로 폭증했습니다. 84㎡ 신고가는 2025년 4월 6.30억까지 오른 뒤 2026년 상반기 5.89억으로 소폭 "
     "내렸고, 평균가도 5.95억→5.67억으로 조정됐습니다. 매물은 탕정과 정반대입니다 — 매매 광고매물이 9주 내내 "
     "550~580건대에서 거의 줄지 않고 정체돼 있습니다(현재 550건, 4개 단지 중 최다). 같은 시기 입주한 신축인데도 "
     "탕정은 물건이 잠기고 성성은 물건이 쌓이는, 상반된 흐름입니다."),
    ("5. 배방지구 — 신축 입주 초기, 매물은 60% 줄고 호가는 실거래보다 앞선다",
     "2025년 준공한 e편한세상탕정퍼스트드림(893세대)은 배방권 순수 아파트 중 최고가입니다. 84㎡ 신고가는 "
     "4.60억(2025년 하반기)→5.50억(2026년 6월)으로 올랐으나 입주 초기라 반기 거래가 2~8건으로 표본이 얇습니다. "
     "매물은 9주간 42건→17건으로 60% 줄었고, 84㎡ 매물 평균 호가(5.90억)는 실거래 신고가(5.50억)보다 7.3% 높습니다. "
     "매도자 눈높이가 실거래를 앞서는, 신축 입주 초기 특유의 모습입니다(다만 표본이 얇아 단정은 이릅니다)."),
    ("6. 매물이 실거래보다 먼저 말한다 — 콕집 데이터가 보여주는 것",
     "실거래는 계약 뒤 30일 내 신고제라 지금 시장의 움직임은 한 달 뒤에야 숫자로 확인됩니다. 반면 매물(호가)은 "
     "매일 바뀌어 시장 심리를 즉각 보여줍니다. 이번 4개 단지에서도 실거래만으로는 드러나지 않는 두 가지가 매물에서 "
     "먼저 나타났습니다. 첫째, ‘매물 잠김 대 쌓임’입니다 — 탕정 지웰시티는 9주간 매물이 –66%로 급감(잠김)한 반면 "
     "성성 레이크사이드는 550건대에서 정체(쌓임)로, 같은 신축권에서도 온도차가 뚜렷합니다. 둘째, ‘광고배율’입니다 — "
     "한 집(실매물)이 여러 중개사무소에 중복 광고된 정도로, 매도 경쟁·회전 강도를 보여주는 콕집 고유 지표입니다. "
     "성성이 3.8배(광고 550건·실매물 144채)로 가장 높고 탕정 3.2배, 불당 2.5배, 배방 2.4배 순이었습니다. "
     "셋째, 호가와 실거래의 간극입니다 — 성숙 단지는 84㎡ 호가가 실거래에 거의 붙어(–2.6~0%) 시장이 정합적인 반면, "
     "신축 배방은 호가가 실거래보다 7.3% 높아 매도자 기대가 앞섰습니다."),
    ("7. 이 자료로 말할 수 있는 것과 없는 것",
     "말씀드릴 수 있는 것은 ‘단계 차이’입니다 — 성숙 단지(불당·탕정)는 신고가·거래량이 견조하고, 신축 입주장"
     "(성성·배방)은 매물이 나오며 가격이 조정되거나 초기 형성되는 중입니다. 말씀드리기 어려운 것은 앞으로의 방향입니다. "
     "실거래는 계약 후 30일 내 신고제라 2026년 6~7월 구간은 아직 대부분 신고되지 않았고(미완성), 신축 두 곳은 표본이 "
     "얇습니다. 또한 이 네 단지는 각 지역의 대장으로, 지역 전체 평균과는 다를 수 있습니다. 분석 방법이나 원자료가 더 "
     "필요하시면 언제든 말씀해 주십시오."),
]

# 원자료 — 대장단지별 (기자 검증용)
# (단지, 소재지, 세대, 준공, 84㎡ 역대신고가, 거래 2024상, 거래 2026상, 현재 매매광고)
RAW = [
    ("천안불당지웰더샵", "천안 서북구 불당동", "685", "2016", "9.46억 (26.4)", "19", "28", "96"),
    ("지웰시티센트럴푸르지오 1·2·3", "아산 탕정면 (탕정지구)", "2,206", "2022", "7.70억 (23.9)", "62", "103", "241"),
    ("천안푸르지오레이크사이드", "천안 서북구 성성동", "1,023", "2023", "6.30억 (25.4)", "3", "34", "550"),
    ("e편한세상탕정퍼스트드림", "아산 배방읍 (배방지구)", "893", "2025", "5.50억 (26.6)", "-", "8", "17"),
]

# 부록 표2 — 매물·호가 (콕집 강점, 2026-07-28)
# (단지, 매매광고, 실매물, 광고배율, 9주 매물변화, 84㎡ 평균호가, 호가-실거래갭)
MAEMUL_RAW = [
    ("천안불당지웰더샵", "96건", "38채", "2.5배", "+5%", "9.21억", "-2.6%"),
    ("지웰시티센트럴푸르지오 1·2·3", "241건", "76채", "3.2배", "-66%", "7.70억 *", "0.0%"),
    ("천안푸르지오레이크사이드", "550건", "144채", "3.8배", "+1%", "6.27억", "-0.5%"),
    ("e편한세상탕정퍼스트드림", "17건", "7채", "2.4배", "-60%", "5.90억", "+7.3%"),
]

NOTE = (
    "작성 | 콕집(koczip.com) 데이터팀 · 런투온라인(대표 황인찬) · runtoonline@gmail.com · 010-5942-8014  "
    "모든 수치는 저희가 자체 구축한 DB에서 2026년 7월 28일 기준으로 집계했습니다. 실거래는 국토부 신고자료(해제거래 제외)"
    "이며 신고 기한 30일로 2026년 상반기 이후 구간은 미완성입니다. 84㎡는 전용면적 80~86㎡ 구간(국민평형)이며, 신고가는 "
    "해당 구간 역대 최고 실거래가입니다. 매매 광고매물은 콕집 일별 수집 기준(포털 노출 광고 건수, 같은 집의 중복 광고 포함)"
    "입니다. 대장아파트는 주상복합을 제외하고 지역 최고가를 선도하는 아파트로, 배방지구는 주상복합(요진와이시티)을 제외한 "
    "순수 아파트 최고가 단지를 골랐습니다. 부록표의 '반기거래'는 해당 반기 6개월간의 전체 평형 거래 건수이며, "
    "약 2년 전(2024년 상반기)과 현재(2026년 상반기)를 비교해 거래량 변화를 보였습니다. "
    "매물·호가 표(표2)에서 광고배율은 광고매물÷실매물, 9주 매물변화는 6월 2일 대비 7월 28일 매매 광고매물 변화율입니다. "
    "84㎡ 평균호가는 해당 평형(공급 111~118㎡) 매물의 평균 호가이며, 호가–실거래갭은 이 평균호가가 84㎡ 역대 신고가 대비 "
    "몇 % 인지를 뜻합니다. 탕정 지웰시티의 84㎡ 호가(*)는 현재 매물이 1건뿐이라 대표성이 낮습니다. 신축(성성 레이크사이드·"
    "배방 e편한세상탕정)은 입주 초기라 표본이 얇은 점을 함께 감안해 주시기 바랍니다."
)


def make_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11.2, 4.1), dpi=200)
    xs = list(range(len(PERIODS)))

    # ① 84㎡ 신고가 추이
    for nm, s in SERIES.items():
        ys = s["hi"]
        xv = [x for x, y in zip(xs, ys) if y is not None]
        yv = [y for y in ys if y is not None]
        ax1.plot(xv, yv, color=s["color"], lw=2.2, marker="o", ms=4.5, label=nm)
        ax1.annotate(f"{yv[-1]:.2f}", (xv[-1], yv[-1]), fontsize=8.2, fontweight="bold",
                     color=s["color"], xytext=(4, 2), textcoords="offset points")
    ax1.set_xticks(xs); ax1.set_xticklabels(PERIODS, fontsize=8.5)
    ax1.set_ylabel("84㎡ 신고가 (억원)", fontsize=9)
    ax1.set_title("① 84㎡ 신고가 추이 — 불당 대장 9.46억 신고가 경신", fontsize=10.3, fontweight="bold", pad=8, loc="left")
    ax1.grid(axis="y", color="#eef2f6", lw=.8)
    ax1.legend(fontsize=7.8, loc="upper left", frameon=False, ncol=1)
    for sp in ("top", "right"):
        ax1.spines[sp].set_visible(False)

    # ② 반기별 거래량
    for nm, s in SERIES.items():
        ys = s["vol"]
        xv = [x for x, y in zip(xs, ys) if y]
        yv = [y for y in ys if y]
        ax2.plot(xv, yv, color=s["color"], lw=2.2, marker="s", ms=4.5, label=nm)
        ax2.annotate(f"{yv[-1]}", (xv[-1], yv[-1]), fontsize=8.2, fontweight="bold",
                     color=s["color"], xytext=(4, 2), textcoords="offset points")
    ax2.set_xticks(xs); ax2.set_xticklabels(PERIODS, fontsize=8.5)
    ax2.set_ylabel("반기 거래량 (건, 전체 평형)", fontsize=9)
    ax2.set_title("② 반기 거래량 — 탕정 지웰시티 103건·성성 입주장 폭증", fontsize=10.3, fontweight="bold", pad=8, loc="left")
    ax2.grid(axis="y", color="#eef2f6", lw=.8)
    for sp in ("top", "right"):
        ax2.spines[sp].set_visible(False)

    fig.suptitle("삼성 탕정 반도체벨트 배후 4개 대장아파트 84㎡ 실거래 (콕집 DB, 2026-07-28)",
                 fontsize=11.3, fontweight="bold", y=0.995)
    fig.text(0.5, 0.005,
             "84㎡=전용 80~86㎡ 구간의 역대 최고 실거래가(①)·반기 전체거래 건수(②). 실거래 신고 30일 지연으로 2026년 상반기 "
             "이후는 미완성이며, 신축(성성·배방)은 입주 초기라 표본이 얇습니다.",
             ha="center", fontsize=7.6, color="#64748b")
    fig.tight_layout(rect=[0, 0.05, 1, 0.94])
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_maemul_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11.2, 4.1), dpi=200)

    # ① 주간 매매 광고매물 추이 (콕집 일별 수집)
    xs = list(range(len(MAEMUL_WEEKS)))
    for nm, ys in MAEMUL_SERIES.items():
        ax1.plot(xs, ys, color=COLORS[nm], lw=2.2, marker="o", ms=3.8, label=nm)
        ax1.annotate(f"{ys[-1]}", (xs[-1], ys[-1]), fontsize=8, fontweight="bold",
                     color=COLORS[nm], xytext=(4, 0), textcoords="offset points", va="center")
    ax1.set_xticks(xs[::2]); ax1.set_xticklabels([MAEMUL_WEEKS[i] for i in xs[::2]], fontsize=8)
    ax1.set_ylabel("매매 광고매물 (건)", fontsize=9)
    ax1.set_title("① 주간 매물 추이 — 탕정 –66% 급감 vs 성성 정체", fontsize=10.3, fontweight="bold", pad=8, loc="left")
    ax1.grid(axis="y", color="#eef2f6", lw=.8)
    ax1.legend(fontsize=7.6, loc="upper right", frameon=False)
    for sp in ("top", "right"):
        ax1.spines[sp].set_visible(False)

    # ② 광고매물 vs 실매물 (광고배율 = 매도 경쟁 강도)
    names = list(MAEMUL_NOW.keys())
    ad = [MAEMUL_NOW[n][0] for n in names]
    unit = [MAEMUL_NOW[n][1] for n in names]
    rat = [MAEMUL_NOW[n][2] for n in names]
    y = list(range(len(names)))[::-1]
    bw = 0.38
    ax2.barh([v + bw / 2 for v in y], ad, bw, color="#9db8dd", label="광고매물(포털 노출)")
    ax2.barh([v - bw / 2 for v in y], unit, bw, color="#1268d3", label="실매물(중복 제거)")
    for yy, a, u, r in zip(y, ad, unit, rat):
        ax2.annotate(f"{a}", (a, yy + bw / 2), va="center", xytext=(3, 0),
                     textcoords="offset points", fontsize=7.6)
        ax2.annotate(f"{u}채 · 광고 {r}배", (max(a, u), yy - bw / 2), va="center", xytext=(3, 0),
                     textcoords="offset points", fontsize=7.6, fontweight="bold", color="#13294b")
    ax2.set_yticks(y); ax2.set_yticklabels([n.split()[0] for n in names], fontsize=9)
    ax2.set_xlim(0, 650)
    ax2.set_xlabel("현재 매매 매물 (건, 2026-07-28)", fontsize=9)
    ax2.set_title("② 광고 vs 실매물 — 광고배율(콕집 고유지표)", fontsize=10.3, fontweight="bold", pad=8, loc="left")
    ax2.legend(fontsize=7.6, loc="lower right", frameon=False)
    for sp in ("top", "right"):
        ax2.spines[sp].set_visible(False)

    fig.suptitle("콕집 매물 데이터 — 실거래(신고 30일 지연)가 못 보여주는 즉각 반응 (2026-07-28)",
                 fontsize=11.3, fontweight="bold", y=0.995)
    fig.text(0.5, 0.005,
             "매물수는 콕집이 매일 수집하는 포털 노출 광고 건수(①②의 '광고매물')이며, 실매물은 같은 집의 중복 광고를 합친 수입니다. "
             "광고배율(광고÷실매물)이 높을수록 한 물건을 여러 중개사무소가 경쟁적으로 광고한 것입니다.",
             ha="center", fontsize=7.6, color="#64748b")
    fig.tight_layout(rect=[0, 0.05, 1, 0.94])
    fig.savefig(CHART2, facecolor="white")
    print("chart2:", CHART2)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"; st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45; st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run(DOCTYPE)
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    t = d.add_paragraph(); r = t.add_run(TITLE); r.font.size = Pt(14); r.font.bold = True
    s2 = d.add_paragraph(); r = s2.add_run(SUBTITLE)
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 요약"); r.font.bold = True; r.font.size = Pt(11.5)
    tb = d.add_table(rows=0, cols=2); tb.style = "Light Grid Accent 1"
    for k, v in SUMMARY:
        c = tb.add_row().cells
        c[0].paragraphs[0].add_run(k).bold = True
        c[1].paragraphs[0].add_run(v)
    for rowx in tb.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.8))

    for head, body in BODY:
        if head.startswith("6."):   # 매물 섹션 앞에 매물 차트 삽입
            ch2 = d.add_paragraph(); ch2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ch2.add_run().add_picture(str(CHART2), width=Cm(16.8))
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)

    def _table(title, heads, rows):
        h = d.add_paragraph(); r = h.add_run(title); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8)
        tt = d.add_table(rows=1, cols=len(heads)); tt.style = "Light Grid Accent 1"
        for c, x in zip(tt.rows[0].cells, heads):
            c.paragraphs[0].add_run(x).bold = True
        for row in rows:
            cells = tt.add_row().cells
            for c, x in zip(cells, row):
                c.paragraphs[0].add_run(x)
        for rowx in tt.rows:
            for c in rowx.cells:
                for pp in c.paragraphs:
                    pp.paragraph_format.space_after = Pt(2)
                    for rr in pp.runs:
                        rr.font.size = Pt(8)

    _table("■ 부록1 — 대장아파트별 84㎡ 실거래 (2026-07-28 기준)",
           ["대장단지", "소재지", "세대", "준공", "84㎡ 역대신고가", "반기거래\n2024상반기", "반기거래\n2026상반기", "현재\n매매광고"], RAW)
    _table("■ 부록2 — 대장아파트별 매물·호가 (콕집 일별 수집, 2026-07-28 기준)",
           ["대장단지", "매매광고", "실매물", "광고배율", "9주 매물변화", "84㎡ 평균호가", "호가–실거래갭"], MAEMUL_RAW)

    m = d.add_paragraph(); r = m.add_run(NOTE)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    d.save(out); print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE); gray = colors.Color(0.39, 0.45, 0.55)
    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=3),
        "title": ParagraphStyle("title", fontName="PB", fontSize=13, leading=18, spaceAfter=3),
        "sub": ParagraphStyle("sub", fontName="P", fontSize=9.2, textColor=gray, leading=13.5, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.7, leading=15.3, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "note": ParagraphStyle("note", fontName="P", fontSize=7.5, leading=11.3, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=7.3, leading=10),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=7.3, leading=10),
        "sumk": ParagraphStyle("sumk", fontName="PB", fontSize=8.6, leading=12.2),
        "sumv": ParagraphStyle("sumv", fontName="P", fontSize=8.6, leading=12.2),
    }
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=15 * mm, rightMargin=15 * mm,
                            topMargin=13 * mm, bottomMargin=13 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart = RLImage(str(CHART), width=180 * mm, height=180 * mm * h / w)
    w2, h2 = PImage.open(CHART2).size
    chart2 = RLImage(str(CHART2), width=180 * mm, height=180 * mm * h2 / w2)

    sum_tbl = Table([[Paragraph(k, ss["sumk"]), Paragraph(v, ss["sumv"])] for k, v in SUMMARY],
                    colWidths=[22 * mm, 158 * mm])
    sum_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.955, 0.97, 0.99)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))

    HEADS = ["대장단지", "소재지", "세대", "준공", "84㎡ 역대신고가",
             "반기거래\n2024상반기", "반기거래\n2026상반기", "현재\n매매광고"]
    tdata = [[Paragraph(x, ss["cellb"]) for x in HEADS]]
    for row in RAW:
        tdata.append([Paragraph(x, ss["cell"]) for x in row])
    _tstyle = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)])
    raw_tbl = Table(tdata, colWidths=[34 * mm, 30 * mm, 14 * mm, 13 * mm, 27 * mm, 14 * mm, 14 * mm, 18 * mm])
    raw_tbl.setStyle(_tstyle)

    HEADS2 = ["대장단지", "매매광고", "실매물", "광고배율", "9주 매물변화", "84㎡ 평균호가", "호가–실거래갭"]
    tdata2 = [[Paragraph(x, ss["cellb"]) for x in HEADS2]]
    for row in MAEMUL_RAW:
        tdata2.append([Paragraph(x, ss["cell"]) for x in row])
    mae_tbl = Table(tdata2, colWidths=[42 * mm, 20 * mm, 18 * mm, 18 * mm, 24 * mm, 24 * mm, 24 * mm])
    mae_tbl.setStyle(_tstyle)

    story = [
        Paragraph(DOCTYPE, ss["tag"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph("■ 요약", ss["head"]), sum_tbl, Spacer(1, 4 * mm),
        chart, Spacer(1, 3 * mm),
    ]
    for head, body in BODY:
        if head.startswith("6."):
            story.append(Spacer(1, 2 * mm))
            story.append(chart2)
            story.append(Spacer(1, 2 * mm))
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph("■ 부록1 — 대장아파트별 84㎡ 실거래 (2026-07-28 기준)", ss["head"]))
    story.append(raw_tbl)
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph("■ 부록2 — 대장아파트별 매물·호가 (콕집 일별 수집, 2026-07-28 기준)", ss["head"]))
    story.append(mae_tbl)
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(NOTE, ss["note"]))
    doc.build(story); print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_maemul_chart()
    make_docx(HERE / "반도체벨트_대장아파트.docx")
    make_pdf(HERE / "반도체벨트_대장아파트.pdf")
