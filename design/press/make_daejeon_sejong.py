# -*- coding: utf-8 -*-
"""대전 선도지구 vs 세종 청사권 매물 추이 보도자료.
수치는 2026-07-20 complex_daily_agg·transactions 실측.
Run: python3 design/press/make_daejeon_sejong.py
 → design/press/대전선도지구_세종매물추이.docx / .pdf + data/daejeon_sejong_chart.png
발송 없음 — 파일 생성만.

※ '세종 투자수요가 대전으로 이동' 가설은 **데이터가 뒷받침하지 않아** 그렇게 쓰지 않았다.
   근거: 발표 후 세종 청사권 매물 -0.5%·호가 +0.19% 로 사실상 무반응. 자금이 빠져나갔다면
   세종에서 매물이 늘거나 호가가 밀려야 하는데 둘 다 없다. 매수자 거주지 자료도 우리에게 없어
   '이동'은 확인 불가. 확인되는 사실(국지적 잠김·격차)만 제목과 본문에 담았다.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent / "fonts"
DATA = HERE / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "daejeon_sejong_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

TITLE = "대전 선도지구는 잠기고, 세종 청사권은 미동 없었다 — 콕집 DB로 본 충청권 두 축"
SUBTITLE = "발표 엿새, 매매 매물 대전 -22.5% vs 세종 -0.5% · 호가는 대전만 8배 더 올라"

LEAD = (
    "7월 15일 대전 노후계획도시 정비 선도지구 발표 이후 대상 단지의 매물이 빠르게 잠긴 반면, "
    "인접한 세종시 정부청사 생활권은 거의 반응하지 않은 것으로 나타났다. 부동산 데이터 플랫폼 "
    "콕집(koczip.com)이 대전 선도지구 6개 단지와 세종 청사권 7개 동 103개 단지의 매물을 일자별로 "
    "전수 비교한 결과, 발표 전날인 7월 14일 대비 7월 20일 매매 광고매물은 대전이 227건에서 176건으로 "
    "22.5% 급감한 반면 세종은 15,385건에서 15,302건으로 0.5% 줄어 사실상 변동이 없었다. "
    "중복 광고를 합친 실매물 기준으로도 대전은 18.8% 줄었지만 세종은 0.2% 감소에 그쳤다."
)

BODY = [
    ("발표 전엔 나란히 늘었다 — 갈라진 건 발표 직후",
     "두 지역이 원래 다른 흐름이었던 것은 아니다. 발표 전 한 달간(6월 18일~7월 14일) 매매 광고매물은 "
     "세종 청사권이 14,786건에서 15,385건으로 4.1%, 대전 선도지구가 217건에서 227건으로 4.6% 늘어 "
     "증가율이 사실상 같았다. 갈라진 시점은 발표 당일이다. 대전은 7월 15일 220건을 기점으로 16일 "
     "206건, 17일 195건, 18일 191건, 19일 187건, 20일 176건으로 엿새 연속 줄었고, 같은 기간 세종은 "
     "15,513→15,302건 사이에서 오르내렸다. 재건축 기대가 실린 곳에서만 매도자가 물건을 거둬들이는 "
     "‘매물 잠김’이 국지적으로 나타난 셈이다."),
    ("세종은 동별로도 평시 등락 — 특정 방향의 쏠림 없어",
     "세종 청사권을 동 단위로 쪼개 봐도 한 방향의 움직임은 확인되지 않는다. 발표 전날 대비 7월 20일 "
     "매매 광고매물은 한솔동이 3,027건에서 2,955건으로 2.4% 줄어 감소 폭이 가장 컸고, 어진동 -1.1%, "
     "종촌동 -1.0%, 다정동 -0.6%, 도담동 -0.1%였다. 반면 새롬동은 0.4%, 나성동은 1.8% 늘었다. "
     "감소와 증가가 뒤섞인 ±2% 안팎의 등락으로, 평상시 하루 단위 변동 범위를 벗어나지 않는다."),
    ("호가는 대전만 올랐다",
     "가격표에서도 온도차가 뚜렷하다. 같은 엿새 동안 단지 전체 평균 매매호가(매물수 가중)는 대전 "
     "선도지구가 8억 3,108만 원에서 8억 4,439만 원으로 1.60% 오른 반면, 세종 청사권은 6억 7,885만 "
     "원에서 6억 8,011만 원으로 0.19% 오르는 데 그쳤다. 상승률로 8배 차이다. 전세 매물은 대전이 "
     "46건에서 54건으로 17.4% 늘고 세종은 773건에서 770건으로 0.4% 줄어, 대전에서 잠긴 것은 "
     "매매뿐이라는 점도 확인된다."),
    ("1년 시계로 보면 — 가격은 벌어지고, 거래는 함께 식었다",
     "시야를 1년으로 넓히면 두 도시의 아파트 평균 실거래가는 반대로 움직였다. 국토교통부 실거래 "
     "신고 기준 아파트 매매 평균가는 대전이 2025년 8월 3억 4,312만 원에서 2026년 6월 3억 6,977만 "
     "원으로 7.8% 오른 반면, 세종은 5억 631만 원에서 5억 698만 원으로 0.1% 움직이는 데 그쳤다. "
     "다만 거래량은 두 도시가 함께 줄었다. 월별 매매 신고 건수는 2026년 1월 세종 589건·대전 1,736건으로 "
     "정점을 찍은 뒤 6월 세종 356건·대전 1,262건으로 각각 39.6%, 27.3% 감소했다. 한쪽이 늘고 "
     "한쪽이 주는 그림은 아니다."),
    ("‘세종 자금이 대전으로 옮겨갔다’고 볼 근거는 확인되지 않았다",
     "이번 분석에서 가장 조심스럽게 다뤄야 할 대목이다. 세종에서 대전으로 투자 수요가 이동했다면 "
     "세종에서 파는 움직임, 즉 매물 증가나 호가 약세가 먼저 나타나야 한다. 그러나 발표 전후 세종 "
     "청사권의 매매 매물은 0.5% 줄었고 호가는 0.19% 올라 어느 쪽 신호도 없었다. 거래량 역시 두 "
     "도시가 같은 시기에 함께 줄었다. 매수자의 거주지를 알 수 있는 자료는 국토교통부 실거래 공개 "
     "항목에 포함되지 않아, 자금의 실제 이동 여부는 이 데이터로 확인할 수 없다. 현재 데이터로 "
     "말할 수 있는 것은 ‘대전 선도지구에서 국지적·이벤트성 매물 잠김이 나타났고, 세종 청사권은 "
     "그 영향을 받지 않았다’는 사실까지다."),
    ("해석 시 유의점",
     "이번 수치는 발표 후 엿새간의 초기 반응이다. 매물 감소에는 호가를 다시 부르기 위한 회수와 실제 "
     "계약 성사가 섞여 있을 수 있다. 발표 이후 계약분의 실거래 신고는 신고 기한이 30일이어서 아직 "
     "집계에 반영되지 않았다. 월별 거래량 비교에서 최근 두 달이 적어 보이는 것도 같은 이유이며, "
     "감소로 읽으면 오독이다. 세종 청사권은 동 단위로 묶은 광역 표본(103개 단지)이고 대전은 발표 "
     "대상 6개 단지여서 표본 성격이 다르다는 점도 함께 고려해야 한다."),
    ("일자별 매물·호가, 누구나 단지 단위로 확인 가능",
     "콕집은 전국 아파트·오피스텔 6만 4천여 단지의 매물과 실거래를 매일 수집해 교차 분석하는 "
     "플랫폼으로, 이번 분석에 쓰인 일자별 매물수·실매물수·평균 호가 추이는 콕집 단지 상세의 "
     "‘매물분석’ 메뉴에서 누구나 무료로 확인할 수 있다. 지역 비교, 급매 탐지, 단지별 신고가 등 "
     "분석 기능도 함께 제공한다."),
]

QUOTE = ("조용호 콕집 공동창업자는 “개발 발표가 나면 인접 도시의 돈이 옮겨갔다는 식의 해석이 "
         "쉽게 붙지만, 데이터로 보면 이번 반응은 대전 선도지구 안에서만 일어났고 세종은 움직이지 "
         "않았다”며 “매수자 거주지처럼 자금 이동을 직접 보여주는 자료가 없는 상태에서 인과를 "
         "단정하기보다, 확인되는 사실과 확인되지 않는 것을 나눠 보여주는 게 소비자에게 도움이 "
         "된다고 본다”고 말했다.")

COMPANY = [
    ("서비스", "콕집 — 부동산 매물·실거래·중개사 분석 플랫폼 (koczip.com)"),
    ("운영사", "런투온라인 (대표 황인찬)"),
    ("문의", "조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946"),
]

METHOD = (
    "분석 방법 | 대상: ①대전 선도지구 3개 구역 6개 단지(한가람·공작한양·목련·크로바·보람·삼익소월) "
    "②세종 정부청사 생활권 7개 동(어진·도담·종촌·다정·나성·새롬·한솔) 103개 단지. 기간: 2026-06-17~07-20 일자별. "
    "광고매물=포털 노출 광고 건수, 실매물=동일 주택의 중복 광고를 합친 수, 평균 호가=매물수 가중평균. "
    "실거래 건수·평균가는 국토교통부 실거래 신고 자료(해제거래 제외) 기준이며 신고 기한 30일로 최근 2개월은 미완성. "
    "호가는 매도 희망가로 실거래 가격이 아님. 전 수치는 콕집 자체 구축 DB 실측(2026-07-20)."
)

# ── 실측 데이터(2026-07-20 DB) ──
DAILY = [  # (날짜, 세종 청사권, 대전 선도지구) — 매매 광고매물
    ("06-17", 15105, 220), ("06-18", 14786, 217), ("06-19", 14971, 221), ("06-20", 15396, 225),
    ("06-21", 15216, 218), ("06-22", 15359, 228), ("06-23", 15406, 226), ("06-24", 15523, 227),
    ("06-25", 15541, 232), ("06-26", 15498, 228), ("06-27", 15545, 228), ("06-28", 15288, 224),
    ("06-29", 15389, 227), ("06-30", 15391, 221), ("07-01", 15536, 226), ("07-02", 15569, 223),
    ("07-03", 15470, 228), ("07-04", 15438, 232), ("07-05", 15049, 229), ("07-06", 15281, 235),
    ("07-07", 15367, 238), ("07-08", 15535, 237), ("07-09", 15414, 235), ("07-10", 15503, 231),
    ("07-11", 15497, 232), ("07-12", 15163, 223), ("07-13", 15348, 231), ("07-14", 15385, 227),
    ("07-15", 15513, 220), ("07-16", 15471, 206), ("07-17", 15332, 195), ("07-18", 15490, 191),
    ("07-19", 15108, 187), ("07-20", 15302, 176),
]

TABLE = [  # 항목, 대전 선도지구, 세종 청사권
    ("매매 광고매물 (7/14→7/20)", "227 → 176", "15,385 → 15,302"),
    ("증감률", "-22.5%", "-0.5%"),
    ("실매물 (중복 합침)", "144 → 117 (-18.8%)", "5,019 → 5,008 (-0.2%)"),
    ("평균 매매호가", "8억 3,108만 → 8억 4,439만 (+1.60%)", "6억 7,885만 → 6억 8,011만 (+0.19%)"),
    ("전세 매물", "46 → 54 (+17.4%)", "773 → 770 (-0.4%)"),
    ("발표 전 한 달 추세 (6/18→7/14)", "217 → 227 (+4.6%)", "14,786 → 15,385 (+4.1%)"),
    ("아파트 평균 실거래가 (25.8→26.6)", "3억 4,312만 → 3억 6,977만 (+7.8%)", "5억 631만 → 5억 698만 (+0.1%)"),
    ("월 매매 신고건수 (26.1→26.6)", "1,736 → 1,262 (-27.3%)", "589 → 356 (-39.6%)"),
]


def make_chart():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    # 규모가 70배 차이(1.5만 vs 227)라 7/14=100 으로 지수화해야 비교가 보인다
    base_i = next(i for i, d in enumerate(DAILY) if d[0] == "07-14")
    bs, bd = DAILY[base_i][1], DAILY[base_i][2]
    xs = list(range(len(DAILY)))
    sj = [d[1] / bs * 100 for d in DAILY]
    dj = [d[2] / bd * 100 for d in DAILY]

    fig, ax = plt.subplots(figsize=(8.6, 3.5), dpi=200)
    ann = next(i for i, d in enumerate(DAILY) if d[0] == "07-15")
    ax.axvline(ann, color="#e2574c", lw=1.2, ls="--", alpha=.9)
    ax.text(ann - 0.4, 108.6, "7/15 대전 선도지구 발표", color="#e2574c", fontsize=9,
            fontweight="bold", ha="right", va="top")
    ax.axhline(100, color="#c9d3de", lw=.9, ls=":")
    ax.plot(xs, sj, color="#7a8ba0", lw=2.0, label="세종 청사권 103개 단지")
    ax.plot(xs, dj, color="#1268d3", lw=2.4, label="대전 선도지구 6개 단지")
    ax.fill_between(xs[ann:], dj[ann:], 100, color="#1268d3", alpha=.08)
    ax.annotate(f"{dj[-1]:.0f}", (xs[-1], dj[-1]), textcoords="offset points", xytext=(5, -3),
                fontsize=9.5, fontweight="bold", color="#1268d3")
    ax.annotate(f"{sj[-1]:.0f}", (xs[-1], sj[-1]), textcoords="offset points", xytext=(5, -3),
                fontsize=9, color="#5f6f80")
    ticks = [i for i, d in enumerate(DAILY) if d[0] in ("06-17", "06-24", "07-01", "07-08", "07-15", "07-20")]
    ax.set_xticks(ticks)
    ax.set_xticklabels([DAILY[i][0].replace("0", "", 1).replace("-", ".") for i in ticks], fontsize=8.5)
    ax.tick_params(axis="y", labelsize=8.5)
    ax.set_ylim(72, 110)
    ax.set_ylabel("7/14 = 100", fontsize=9)
    ax.grid(axis="y", color="#e8edf3", lw=.7)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    ax.legend(fontsize=9, frameon=False, loc="lower left")
    ax.set_title("매매 광고매물 추이 — 발표 전날(7/14) 대비 지수 (2026.6.17~7.20, 콕집 DB)",
                 fontsize=10.5, fontweight="bold", pad=8)
    fig.tight_layout()
    fig.savefig(CHART, facecolor="white")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run("보도자료")
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)
    p2 = d.add_paragraph()
    r = p2.add_run("배포일: 2026년 7월 · 즉시 보도 가능    문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(15.5); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(11.5); r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3); r.font.bold = True

    d.add_paragraph(LEAD)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.2))

    h = d.add_paragraph(); r = h.add_run("■ 한눈에 보는 비교 (콕집 DB)")
    r.font.bold = True; r.font.size = Pt(11.5)
    tbl = d.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ["항목", "대전 선도지구 (6단지)", "세종 청사권 (103단지)"]):
        c.paragraphs[0].add_run(txt).bold = True
    for row in TABLE:
        for c, txt in zip(tbl.add_row().cells, row):
            c.paragraphs[0].add_run(txt)
    for rowx in tbl.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9)

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)
    q = d.add_paragraph(); r = q.add_run(QUOTE); r.font.italic = True

    m = d.add_paragraph(); r = m.add_run(METHOD)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 서비스 개요"); r.font.bold = True; r.font.size = Pt(11.5)
    for k, v in COMPANY:
        d.add_paragraph(f"· {k}: {v}")
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=2),
        "meta": ParagraphStyle("meta", fontName="P", fontSize=8, textColor=gray, spaceAfter=10),
        "title": ParagraphStyle("title", fontName="PB", fontSize=15, leading=21, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="PB", fontSize=10.8, textColor=blue, leading=15, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "quote": ParagraphStyle("quote", fontName="P", fontSize=9.8, leading=15.5,
                                leftIndent=8, textColor=colors.Color(0.2, 0.25, 0.33), spaceAfter=10),
        "method": ParagraphStyle("method", fontName="P", fontSize=7.8, leading=11.5, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=8.4, leading=11.5),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=8.4, leading=11.5),
    }

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 174 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    def cell(txt, bold=False):
        return Paragraph(txt, ss["cellb"] if bold else ss["cell"])

    tdata = [[cell(x, True) for x in ["항목", "대전 선도지구 (6단지)", "세종 청사권 (103단지)"]]]
    for row in TABLE:
        tdata.append([cell(row[0], True), cell(row[1]), cell(row[2])])
    tbl = Table(tdata, colWidths=[52 * mm, 60 * mm, 62 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))

    story = [
        Paragraph("보도자료", ss["tag"]),
        Paragraph("배포일: 2026년 7월 · 즉시 보도 가능 &nbsp;&nbsp; 문의: 조용호 공동창업자 · albooooza@gmail.com · 010-4692-4946", ss["meta"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph(LEAD, ss["body"]),
        chart,
        Spacer(1, 4 * mm),
        Paragraph("■ 한눈에 보는 비교 (콕집 DB)", ss["head"]),
        tbl,
        Spacer(1, 2 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph(QUOTE, ss["quote"]))
    story.append(Paragraph(METHOD, ss["method"]))
    story.append(Paragraph("■ 서비스 개요", ss["head"]))
    for k, v in COMPANY:
        story.append(Paragraph(f"· <b>{k}</b>: {v}", ss["body"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "대전선도지구_세종매물추이.docx")
    make_pdf(HERE / "대전선도지구_세종매물추이.pdf")
