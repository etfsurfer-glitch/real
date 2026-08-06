# -*- coding: utf-8 -*-
"""매도인 자금 제공 매물 — 기자 전달용 데이터 분석 보고서.
수치는 2026-07-21 listings_current 실측(설명 보유 1,754,823건 전수).
Run: python3 design/press/매도인자금/make_seller_mortgage.py
 → design/press/매도인자금/매도인자금제공_분석보고서.docx / .pdf + data/seller_mortgage_chart.png
발송 없음 — 파일 생성만.

※ 보도자료가 아니라 **보고서**다. 차이를 지킨다:
  - 대표 인용문·서비스 홍보 없음. 기사 문장을 대신 써주지 않는다.
  - 원자료 6건을 표로 전부 공개해 기자가 직접 검증할 수 있게 한다.
  - '말할 수 없는 것'과 '추가 취재가 필요한 지점'을 결과만큼 비중 있게 쓴다.
※ 정치적으로 민감한 사안이라 지킬 선:
  1) 이 대통령 거래의 성격(편법이냐 정상이냐)은 평가하지 않는다. 우리 데이터로 알 수 없다.
  2) '전국에 6건뿐인 희귀 거래'로 읽히면 안 된다. 광고 문구 기준이지 등기부 기준이 아니다.
  3) 6건은 **광고 건수**이지 물건 수가 아니다(레이크팰리스 3건이 동일 면적·동일 호가).
  4) 어느 정당의 주장도 인용하거나 편들지 않는다.
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent
FONTS = HERE.parent.parent / "fonts"
DATA = HERE.parent / "data"
DATA.mkdir(exist_ok=True)
CHART = DATA / "seller_mortgage_chart.png"
BLUE = (0x12 / 255, 0x68 / 255, 0xD3 / 255)

DOCTYPE = "데이터 분석 보고서 (취재 참고자료)"
TITLE = "매도인이 매수인에게 직접 자금을 대주는 매물, 전국 광고 175만 건 중 6건"
SUBTITLE = "전국 아파트·오피스텔 매물 광고 설명 문구 전수 분석 · 2026년 7월 21일 기준"

SUMMARY = [
    ("분석 대상", "전국 아파트·오피스텔 매물 광고 중 설명 문구가 있는 1,754,823건 전수"),
    ("집계 결과", "매도인이 매수인에게 직접 자금을 대준다고 밝힌 광고 6건 (단지 기준 4곳)"),
    ("대조군", "‘근저당 없음·무융자’ 등 담보가 없다는 점을 내세운 광고 8,167건 (1,361배)"),
    ("공통점", "6건 모두 매매 호가 20억원 이상 · 서울 강남권·한강변과 경기 분당"),
    ("유의 사항", "광고 문구를 집계한 수치이며, 등기부상 실제 근저당 설정 건수가 아닙니다"),
]

BODY = [
    ("1. 분석 배경",
     "부동산 매매에서 매도인이 매수인 앞으로 근저당권을 설정하는 방식, 즉 매도인이 매매대금의 "
     "일부를 사실상 빌려주고 그 채권을 담보로 잡는 거래 형태가 최근 논의되고 있습니다. 다만 이 "
     "방식이 시장에서 얼마나 통용되는지에 대한 공개 통계는 확인되지 않았습니다. 이에 저희가 "
     "자체 수집하는 전국 매물 광고 데이터에서 중개사가 이 조건을 명시해 광고한 사례가 얼마나 "
     "되는지를 집계해 보았습니다. 본 보고서는 그 집계 결과와 한계를 정리한 것으로, 특정 거래의 "
     "적법성이나 통상성에 대한 평가는 담고 있지 않습니다."),
    ("2. 분석 대상과 방법",
     "저희가 매일 수집하는 전국 아파트·오피스텔 매물 광고 가운데 설명 문구(중개사가 직접 작성하는 "
     "매물 소개란)가 있는 1,754,823건을 대상으로 했습니다. 설명 문구에서 ‘매도인·집주인·소유자’ 등 "
     "주체를 나타내는 말과 ‘대출·융자·근저당·담보·대여’ 등 자금을 나타내는 말이 붙어 있거나 "
     "인접한 경우를 추출한 뒤, 전수를 육안으로 확인해 분류했습니다. "
     "분류 과정에서 다음 세 가지는 제외했습니다. 첫째, ‘근저당없음·무근저당·주인대출없음·집주인대출X·"
     "무융자’ 등 반대 의미의 표기입니다. 주체어와 자금어가 함께 나온 29건 중 24건(83%)이 여기 "
     "해당했습니다. 둘째, ‘대출 6억 가능’처럼 매수인이 금융기관에서 받을 수 있는 한도를 뜻하는 "
     "문구입니다. 셋째, ‘기존대출승계·담보대출승계’와 분양권·입주권의 중도금·분담금 승계입니다. "
     "이는 매수인이 매도인의 기존 금융기관 대출을 넘겨받는 것이어서 매도인이 자금을 대주는 것과는 "
     "다르다고 판단했습니다. ‘잔금유예’ 411건도 대부분 시행사의 신축·미분양 판매 조건이어서 "
     "제외했습니다."),
    ("3. 결과 (1) — 해당 광고는 6건",
     "매도인이 매수인에게 직접 자금을 대준다고 밝힌 광고는 6건으로 집계됐습니다. 표기는 "
     "‘매도인근저당가능’, ‘매도인 대출 가능’, ‘매도인 잔금5억 대여가능’, ‘집주인대출6억’, "
     "‘집주인대출 7억 가능’이었습니다. 6건 전체의 원문과 소재지는 아래 부록 표에 그대로 옮겨 "
     "두었으니 직접 확인해 보시기 바랍니다. "
     "다만 이 숫자는 광고 건수이며 물건 수가 아니라는 점을 함께 봐주시면 좋겠습니다. 서울 송파구 "
     "레이크팰리스 3건은 동일 면적(86A)에 동일 호가(28억6,000만원)여서, 같은 물건이 복수 "
     "중개사무소를 통해 광고됐을 가능성이 있습니다. 단지 기준으로는 4곳이며, 실제 물건 수는 "
     "4건 이하일 수 있습니다."),
    ("3. 결과 (2) — 반대 문구는 8,167건",
     "같은 데이터에서 ‘근저당 없음’, ‘무근저당’, ‘무융자’, ‘주인대출없음’처럼 담보가 없다는 점을 "
     "내세운 광고는 8,167건이었습니다. 매도인 자금 제공 6건의 1,361배입니다. ‘근저당’이라는 단어가 "
     "들어간 광고 408건만 놓고 보면 270건(66%)이 ‘없음·무·해지’ 같은 부정형이었습니다. 매물 "
     "광고에서 근저당은 내세우는 조건이라기보다 없다는 점을 강조하는 항목으로 쓰이고 있는 것으로 "
     "보입니다."),
    ("3. 결과 (3) — 유사 조건과의 비교",
     "매수인의 초기 자금 부담을 줄이는 조건 자체가 드문 것은 아닙니다. 기존 임차인의 보증금을 "
     "떠안고 매수하는 ‘세안고’ 매물은 148,942건으로 전체 광고의 8.5%였습니다. 매도인이 매도 후 "
     "그 집에 전세로 계속 거주하는 ‘주인전세’ 조건은 726건이었습니다. 자기자본을 줄이는 거래 조건은 "
     "임대차를 활용하는 형태가 광고 시장에서 일반적이며, 매도인이 직접 자금을 대주는 형태는 이에 "
     "비해 드물게 나타났습니다."),
    ("4. 이 자료의 한계",
     "본 자료로 말씀드릴 수 있는 것은 ‘매물을 광고하는 단계에서 이 조건을 내거는 사례가 얼마나 "
     "되는가’ 한 가지입니다. 말씀드리기 어려운 부분이 더 많아 미리 밝혀 둡니다. "
     "첫째, 이 수치는 등기부상 실제 근저당 설정 건수가 아닙니다. 매도인 근저당은 매물을 내놓는 "
     "단계가 아니라 매수인과의 계약 협상 과정에서 정해지는 경우가 많고, 그 경우 광고 문구에는 "
     "남지 않습니다. 실제 빈도는 등기 자료로만 확인할 수 있는데 저희는 그 자료를 보유하고 있지 "
     "않습니다. 따라서 ‘전국에서 이런 거래가 6건 있었다’는 의미로는 읽히기 어려운 수치입니다. "
     "둘째, 6건은 표본이라고 하기 어려운 규모입니다. 6건이 모두 20억원 이상이라는 공통점이 "
     "관찰되지만, 이를 ‘고가 구간에서 나타나는 경향’으로 일반화하기에는 근거가 부족합니다. "
     "셋째, 특정 거래가 통상적인지 이례적인지는 본 자료로 판단할 수 없습니다. 광고 시장의 분포와 "
     "실제 거래 관행은 서로 다른 층위의 문제이기 때문입니다."),
    ("5. 저희가 답변드리기 어려운 영역",
     "취재 과정에서 아래 사항을 문의하실 수 있을 듯하여, 저희가 자료를 갖고 있지 않은 부분을 "
     "미리 말씀드립니다. "
     "① 등기부 기준 실제 설정 빈도 — 매매를 원인으로 한 소유권이전 등기와 동시에 매도인을 "
     "근저당권자로 하는 설정 등기가 얼마나 발생하는지는 등기 자료의 영역이어서 저희 데이터로는 "
     "확인되지 않습니다. "
     "② 세무 처리 — 매도인이 자금을 대여할 경우의 이자소득 신고나 증여 추정 여부는 저희가 "
     "판단할 수 있는 사안이 아닙니다. "
     "③ 중개 실무 — 이 조건이 계약서에 어떻게 반영되는지는 저희가 확인하지 못한 부분입니다. "
     "이 세 가지는 본 보고서의 수치만으로는 답변드릴 수 없는 영역임을 양해해 주시기 바랍니다."),
]

# 원자료 — 기자가 직접 확인할 수 있게 전부 공개
RAW = [
    ("포제스한강", "서울 광진구 광장동", "54.0억", "168㎡",
     "매도인근저당가능 방음벽너머 한강뷰 개방감최고", "07-13"),
    ("레이크팰리스", "서울 송파구 잠실동", "28.6억", "86A",
     "로얄동,잠실역 초역세권,올수리,매도인 대출 가능,전세끼고 매매", "07-20"),
    ("레이크팰리스", "서울 송파구 잠실동", "28.6억", "86A",
     "로얄동,잠실역 초역세권,올수리,매도인 대출 가능,전세 끼고 매매", "07-16"),
    ("레이크팰리스", "서울 송파구 잠실동", "28.6억", "86A",
     "26P,로얄동,조망굿,전세끼고,확장형,깨끗함,매도인 잔금5억 대여가능", "07-15"),
    ("판교원4단지휴먼시아푸르지오", "경기 성남 분당구 판교동", "23.0억", "186㎡",
     "단독주택 같은 분위기, 앞에 마당, 집주인대출 7억 가능, 서판교역 호재", "07-10"),
    ("강남자곡아이파크", "서울 강남구 자곡동", "20.0억", "81A",
     "집주인대출6억. 수서역세권 개발호재. 추천매물. 자곡초. 풍문고 학세권.", "06-29"),
]

NOTE = (
    "작성 | 콕집(koczip.com) 데이터팀 · 런투온라인(대표 황인찬) · runtoonline@gmail.com · 010-5942-8014  "
    "모든 수치는 저희가 자체 구축한 DB에서 2026년 7월 21일 기준으로 집계했습니다. 부록의 광고 "
    "원문은 중개사가 작성한 문구를 그대로 옮긴 것이며, 확인일자는 중개사가 매물을 최종 확인한 "
    "날짜입니다. 광고는 수시로 등록·삭제되므로 열람 시점에 따라 건수가 달라질 수 있습니다. "
    "분석 방법이나 원자료에 대해 더 확인이 필요하시면 언제든 연락 주시기 바랍니다. 추가 집계나 "
    "다른 조건의 교차 분석이 필요하시면 요청하시는 대로 도와드리겠습니다."
)

# ── 실측(2026-07-21) ──
NEG_ADS = 8167
SELLER_FIN = 6
TENANT = 148942
OWNER = 726
TOTAL = 1754823


# 매매 광고의 호가 구간별 구성 — (5억 미만, 5~10억, 10~20억, 20억 이상, 총건수)
# 건수만 그리면 '6건은 작다'로 끝나지만, 구성비를 겹쳐 놓으면 그 6건이 시장의 어느 자리에
# 있는지가 보인다. 전체 매매를 대조군으로 함께 깔아야 '20억 이상 100%'가 얼마나 치우친
# 값인지 판단할 수 있다(전체 시장에서 20억 이상은 6.1%다).
PRICE_MIX = [
    ("전체 매매 광고 (대조군)", [760206, 497286, 156194, 92478], 1506164),
    ("세안고 (기존 임차인 승계)", [80290, 48587, 13898, 6126], 148901),
    ("주인전세 (매도인이 전세 거주)", [227, 315, 179, 4], 725),
    ("매도인이 직접 자금 제공", [0, 0, 0, 6], 6),
]
BANDS = ["5억 미만", "5~10억", "10~20억", "20억 이상"]
BAND_COLORS = ["#dbe4ee", "#a9bed6", "#5b86bd", "#d64545"]


def make_chart():
    """핵심 한 장 — 조건별 호가 구간 구성비(100% 누적) + 오른쪽에 실제 광고 수.

    건수 비교(로그 눈금)는 뺐다. '6건 대 14만 건'은 숫자를 읽으면 끝나는 정보라
    그림이 보탤 게 없고, 로그 축은 오해만 부른다. 대신 '그 6건이 어느 가격대에
    있는가'를 보이면 대조군과 비교해 판단할 거리가 생긴다.
    n=6은 구성비를 그리기엔 너무 작은 표본이라, 막대 끝 건수와 캡션에 못박아 둔다.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    for f in ("Pretendard-Regular.ttf", "Pretendard-Bold.ttf"):
        font_manager.fontManager.addfont(str(FONTS / f))
    plt.rcParams["font.family"] = "Pretendard"
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(10.4, 3.6), dpi=200)
    ys = list(range(len(PRICE_MIX)))[::-1]

    for yi, (label, counts, n) in zip(ys, PRICE_MIX):
        left = 0.0
        for ci, (cnt, col) in enumerate(zip(counts, BAND_COLORS)):
            pct = 100.0 * cnt / n
            if pct <= 0:
                continue
            ax.barh(yi, pct, 0.58, left=left, color=col,
                    label=BANDS[ci] if yi == ys[0] else None,
                    edgecolor="white", linewidth=0.8)
            # 좁은 조각에 글씨를 넣으면 뭉개진다 — 6% 이상만 안에 표기
            if pct >= 6:
                ax.annotate(f"{pct:.0f}%", (left + pct / 2, yi), ha="center", va="center",
                            fontsize=8.6, fontweight="bold",
                            color="white" if ci >= 2 else "#334155")
            left += pct
        ax.annotate(f"광고 {n:,}건", (101, yi), va="center", ha="left",
                    fontsize=8.8, fontweight="bold", color="#334155")

    ax.set_yticks(ys)
    ax.set_yticklabels([p[0] for p in PRICE_MIX], fontsize=9.6)
    ax.set_xlim(0, 118)
    ax.set_xticks([0, 25, 50, 75, 100])
    ax.set_xticklabels(["0", "25", "50", "75", "100%"], fontsize=8.5)
    ax.set_xlabel("호가 구간별 구성비", fontsize=9)
    ax.grid(axis="x", color="#eef2f6", lw=.8)
    for sp in ("top", "right", "left"):
        ax.spines[sp].set_visible(False)
    ax.legend(ncol=4, fontsize=8.6, frameon=False,
              loc="lower center", bbox_to_anchor=(0.5, 1.02))

    fig.suptitle("매도인이 직접 자금을 대주는 매물은 전부 20억 이상 — 전체 시장에선 20억 이상이 6%",
                 fontsize=11.2, fontweight="bold", y=1.06)
    fig.text(0.5, -0.02,
             "전국 매매 광고 1,506,164건 중 호가가 있는 건 기준(콕집 DB, 2026-07-21). "
             "맨 아래 막대는 6건이어서 구성비 100%에 통계적 의미를 두기는 어렵습니다 — 6건이 모두 20억 이상이었다는 사실의 표시로만 봐주시기 바랍니다. "
             "광고 문구를 집계한 수치이며 등기부상 실제 설정 건수가 아닙니다.",
             ha="center", fontsize=7.8, color="#64748b")
    fig.tight_layout(rect=[0, 0.02, 1, 0.93])
    fig.savefig(CHART, facecolor="white", bbox_inches="tight")
    print("chart:", CHART)


def make_docx(out: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    st = d.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.paragraph_format.line_spacing = 1.45
    st.paragraph_format.space_after = Pt(10)

    p = d.add_paragraph(); r = p.add_run(DOCTYPE)
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor(0x12, 0x68, 0xD3)

    t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(TITLE); r.font.size = Pt(14.5); r.font.bold = True
    s = d.add_paragraph(); r = s.add_run(SUBTITLE)
    r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    h = d.add_paragraph(); r = h.add_run("■ 요약"); r.font.bold = True; r.font.size = Pt(11.5)
    tb = d.add_table(rows=0, cols=2); tb.style = "Light Grid Accent 1"
    for k, v in SUMMARY:
        c = tb.add_row().cells
        c[0].paragraphs[0].add_run(k).bold = True
        c[1].paragraphs[0].add_run(v)
    for rowx in tb.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(9)

    ch = d.add_paragraph(); ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.add_run().add_picture(str(CHART), width=Cm(16.4))

    for head, body in BODY:
        h = d.add_paragraph(); r = h.add_run("■ " + head); r.font.bold = True; r.font.size = Pt(11.5)
        h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(4)
        d.add_paragraph(body)

    h = d.add_paragraph(); r = h.add_run("■ 부록 — 해당 광고 6건 원문 전체")
    r.font.bold = True; r.font.size = Pt(11.5)
    HEADS = ["단지", "소재지", "호가", "면적", "광고 원문", "확인일"]
    tb2 = d.add_table(rows=1, cols=len(HEADS)); tb2.style = "Light Grid Accent 1"
    for c, x in zip(tb2.rows[0].cells, HEADS):
        c.paragraphs[0].add_run(x).bold = True
    for row in RAW:
        for c, x in zip(tb2.add_row().cells, row):
            c.paragraphs[0].add_run(x)
    for rowx in tb2.rows:
        for c in rowx.cells:
            for pp in c.paragraphs:
                pp.paragraph_format.space_after = Pt(2)
                for rr in pp.runs:
                    rr.font.size = Pt(8)

    m = d.add_paragraph(); r = m.add_run(NOTE)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    d.save(out)
    print("docx:", out)


def make_pdf(out: Path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                    Table, TableStyle)
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("P", str(FONTS / "Pretendard-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("PB", str(FONTS / "Pretendard-Bold.ttf")))
    blue = colors.Color(*BLUE)
    gray = colors.Color(0.39, 0.45, 0.55)

    ss = {
        "tag": ParagraphStyle("tag", fontName="PB", fontSize=10, textColor=blue, spaceAfter=3),
        "title": ParagraphStyle("title", fontName="PB", fontSize=14, leading=19, spaceAfter=3),
        "sub": ParagraphStyle("sub", fontName="P", fontSize=9.6, textColor=gray, leading=14, spaceAfter=11),
        "body": ParagraphStyle("body", fontName="P", fontSize=9.8, leading=15.5, spaceAfter=8),
        "head": ParagraphStyle("head", fontName="PB", fontSize=11, leading=15, spaceBefore=6, spaceAfter=4),
        "note": ParagraphStyle("note", fontName="P", fontSize=7.6, leading=11.5, textColor=gray, spaceAfter=8),
        "cell": ParagraphStyle("cell", fontName="P", fontSize=7.6, leading=10.5),
        "cellb": ParagraphStyle("cellb", fontName="PB", fontSize=7.6, leading=10.5),
        "sumk": ParagraphStyle("sumk", fontName="PB", fontSize=8.8, leading=12.5),
        "sumv": ParagraphStyle("sumv", fontName="P", fontSize=8.8, leading=12.5),
    }
    GRID = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.91, 0.945, 0.988)),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ])

    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=16 * mm, rightMargin=16 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)
    from PIL import Image as PImage
    w, h = PImage.open(CHART).size
    chart_w = 178 * mm
    chart = RLImage(str(CHART), width=chart_w, height=chart_w * h / w)

    sum_tbl = Table([[Paragraph(k, ss["sumk"]), Paragraph(v, ss["sumv"])] for k, v in SUMMARY],
                    colWidths=[26 * mm, 152 * mm])
    sum_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.85, 0.88, 0.92)),
        ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.955, 0.97, 0.99)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))

    HEADS = ["단지", "소재지", "호가", "면적", "광고 원문", "확인일"]
    raw_tbl = Table(
        [[Paragraph(x, ss["cellb"]) for x in HEADS]] +
        [[Paragraph(x, ss["cell"]) for x in row] for row in RAW],
        colWidths=[30 * mm, 30 * mm, 14 * mm, 13 * mm, 78 * mm, 13 * mm])
    raw_tbl.setStyle(GRID)

    story = [
        Paragraph(DOCTYPE, ss["tag"]),
        Paragraph(TITLE, ss["title"]),
        Paragraph(SUBTITLE, ss["sub"]),
        Paragraph("■ 요약", ss["head"]), sum_tbl, Spacer(1, 4 * mm),
        chart, Spacer(1, 3 * mm),
    ]
    for head, body in BODY:
        story.append(Paragraph("■ " + head, ss["head"]))
        story.append(Paragraph(body, ss["body"]))
    story.append(Paragraph("■ 부록 — 해당 광고 6건 원문 전체", ss["head"]))
    story.append(raw_tbl)
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(NOTE, ss["note"]))
    doc.build(story)
    print("pdf:", out)


if __name__ == "__main__":
    make_chart()
    make_docx(HERE / "매도인자금제공_분석보고서.docx")
    make_pdf(HERE / "매도인자금제공_분석보고서.pdf")
